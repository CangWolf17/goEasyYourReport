from __future__ import annotations

import json
import shutil
import subprocess
import unittest
import uuid
from pathlib import Path


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PYTHON = Path(r"D:\Miniconda\python.exe")


class ProjectHarness(unittest.TestCase):
    def create_project(self) -> Path:
        sandbox_root = PROJECT_ROOT / "temp" / "equation-bibliography-tests"
        sandbox_root.mkdir(parents=True, exist_ok=True)
        project_root = sandbox_root / uuid.uuid4().hex
        project_root.mkdir(parents=True, exist_ok=True)
        self.addCleanup(lambda: shutil.rmtree(project_root, ignore_errors=True))
        result = subprocess.run(
            [
                str(PYTHON),
                str(PROJECT_ROOT / "scripts" / "init_project.py"),
                "--project-root",
                str(project_root),
            ],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        return project_root

    def run_completed(
        self, project_root: Path, script_name: str, *extra_args: str
    ) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [
                str(PYTHON),
                str(project_root / "scripts" / script_name),
                "--project-root",
                str(project_root),
                *extra_args,
            ],
            capture_output=True,
            text=True,
        )

    def run_workflow(
        self, project_root: Path, action: str, *extra_args: str
    ) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [
                str(PYTHON),
                str(project_root / "scripts" / "workflow_agent.py"),
                action,
                "--project-root",
                str(project_root),
                *extra_args,
            ],
            capture_output=True,
            text=True,
        )

    def write_markdown(self, content: str) -> Path:
        sandbox_root = PROJECT_ROOT / "temp" / "equation-bibliography-tests"
        sandbox_root.mkdir(parents=True, exist_ok=True)
        markdown_path = sandbox_root / f"{uuid.uuid4().hex}.md"
        markdown_path.write_text(content, encoding="utf-8")
        self.addCleanup(lambda: markdown_path.unlink(missing_ok=True))
        return markdown_path


class BibliographyStrategyTests(ProjectHarness):
    def write_bibliography_sources(
        self, project_root: Path, entries: list[dict[str, object]]
    ) -> None:
        source_path = project_root / "logs" / "bibliography.sources.json"
        source_path.parent.mkdir(parents=True, exist_ok=True)
        source_path.write_text(
            json.dumps(entries, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

    def update_bibliography_plan(
        self,
        project_root: Path,
        *,
        source_mode: str,
        output_block_present: bool,
    ) -> None:
        plan_path = project_root / "config" / "template.plan.json"
        plan = json.loads(plan_path.read_text(encoding="utf-8"))
        bibliography = plan["semantics"]["bibliography"]
        bibliography["source_mode"] = source_mode
        bibliography["output_block_present"] = output_block_present
        plan_path.write_text(
            json.dumps(plan, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

    def test_init_project_seeds_bibliography_source_preference_placeholder(self) -> None:
        project_root = self.create_project()

        user_profile = (project_root / "user" / "user.md").read_text(encoding="utf-8")

        self.assertIn("- 参考文献来源：needs_confirmation", user_profile)
        self.assertIn("agent_generate_verified_only", user_profile)
        self.assertIn("agent_search_and_screen", user_profile)
        self.assertIn("user_supplied_files", user_profile)

    def test_preview_summary_surfaces_bibliography_source_confirmation(self) -> None:
        project_root = self.create_project()

        result = self.run_completed(project_root, "build_preview.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        summary = json.loads((project_root / "out" / "preview.summary.json").read_text(encoding="utf-8"))

        self.assertEqual(
            summary["semantics"]["bibliography"]["source_mode"], "needs_confirmation"
        )
        self.assertTrue(summary["semantics"]["bibliography"]["output_block_present"])
        self.assertIn(
            "confirm bibliography source mode",
            summary["review"]["needs_confirmation"],
        )

    def test_workflow_agent_prepare_reports_bibliography_source_decision_needed(
        self,
    ) -> None:
        project_root = self.create_project()

        result = self.run_workflow(project_root, "prepare")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        payload = json.loads(result.stdout)

        self.assertEqual(payload["status"], "ok")
        self.assertTrue(
            any(
                issue.get("kind") == "decision_required"
                and issue.get("details") == "confirm bibliography source mode"
                for issue in payload["issues"]
            )
        )

    def test_bibliography_is_source_only_when_output_block_missing(self) -> None:
        from scripts._bibliography import normalize_bibliography_entries, should_emit_bibliography

        blocks = [
            {"kind": "heading", "level": 2, "text": "参考文献"},
            {"kind": "paragraph", "text": "Author. Title[J]. Journal, 2024."},
        ]
        plan = {"semantics": {"bibliography": {"output_block_present": False}}}

        entries = normalize_bibliography_entries(blocks)

        self.assertEqual(len(entries), 1)
        self.assertFalse(should_emit_bibliography(plan))

    def test_bibliography_defaults_to_bracket_numbered_paragraphs(self) -> None:
        from scripts._bibliography import normalize_bibliography_entries

        blocks = [
            {"kind": "heading", "level": 2, "text": "参考文献"},
            {"kind": "paragraph", "text": "Author. Title[J]. Journal, 2024."},
            {"kind": "paragraph", "text": "Author. Book[M]. Press, 2023."},
        ]

        entries = normalize_bibliography_entries(blocks)

        self.assertEqual(entries[0]["visible_label"], "[1]")
        self.assertEqual(entries[0]["rendered_text"], "[1] Author. Title[J]. Journal, 2024.")
        self.assertEqual(entries[1]["visible_label"], "[2]")
        self.assertEqual(entries[1]["rendered_text"], "[2] Author. Book[M]. Press, 2023.")

    def test_bibliography_registry_reserves_clickable_bracket_label(self) -> None:
        from scripts._bibliography import normalize_bibliography_entries

        blocks = [
            {"kind": "heading", "level": 2, "text": "参考文献"},
            {"kind": "paragraph", "text": "Author. Title[J]. Journal, 2024."},
        ]

        entries = normalize_bibliography_entries(blocks)

        self.assertEqual(entries[0]["id"], "ref_0001")
        self.assertEqual(entries[0]["bookmark"], "ref_0001")
        self.assertEqual(entries[0]["visible_label"], "[1]")

    def test_bibliography_cross_reference_links_whole_bracket_token(self) -> None:
        project_root = self.create_project()
        self.update_bibliography_plan(
            project_root,
            source_mode="agent_generate_verified_only",
            output_block_present=True,
        )
        self.write_bibliography_sources(
            project_root,
            [
                {
                    "id": "ref_0001",
                    "title": "Example Paper",
                    "authors": ["Alice", "Bob"],
                    "year": "2024",
                    "container": "Journal Name",
                    "doi": "10.1000/example",
                    "evidence": {
                        "source_mode": "agent_generate_verified_only",
                        "verified_by_agent": True,
                    },
                }
            ],
        )
        (project_root / "docs" / "report_body.md").write_text(
            "[[REF:bibliography:ref_0001|]]\n\n## 参考文献\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import docx

        redacted = docx.Document(project_root / "out" / "redacted.docx")
        reference_paragraph = next(item for item in redacted.paragraphs if item.text.strip() == "[1]")

        self.assertIn("w:hyperlink", reference_paragraph._p.xml)
        self.assertIn('w:anchor="ref_0001"', reference_paragraph._p.xml)
        self.assertIn(">[1]<", reference_paragraph._p.xml)

    def test_generated_bibliography_entries_use_body_font_and_no_hanging_indent(
        self,
    ) -> None:
        project_root = self.create_project()
        self.update_bibliography_plan(
            project_root,
            source_mode="agent_generate_verified_only",
            output_block_present=True,
        )
        self.write_bibliography_sources(
            project_root,
            [
                {
                    "id": "ref_0001",
                    "title": "Example Paper",
                    "authors": ["Alice", "Bob"],
                    "year": "2024",
                    "container": "Journal Name",
                    "doi": "10.1000/example",
                }
            ],
        )
        (project_root / "docs" / "report_body.md").write_text(
            "## 参考文献\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import docx

        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(item for item in redacted.paragraphs if item.text.strip().startswith("[1]"))

        self.assertEqual(paragraph.style.name, "参考文献")
        self.assertEqual(paragraph.paragraph_format.left_indent.pt, 0.0)
        self.assertEqual(paragraph.paragraph_format.first_line_indent.pt, 0.0)
        self.assertEqual(paragraph.runs[0].font.size.pt, 10.5)

    def test_reference_section_list_entries_reset_to_flush_left_format(self) -> None:
        project_root = self.create_project()
        self.update_bibliography_plan(
            project_root,
            source_mode="agent_generate_verified_only",
            output_block_present=True,
        )
        (project_root / "docs" / "report_body.md").write_text(
            "## 参考文献\n\n1. Author. Title[J]. Journal, 2024.\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import docx

        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(item for item in redacted.paragraphs if "Author. Title" in item.text)

        self.assertEqual(paragraph.style.name, "参考文献")
        self.assertEqual(paragraph.paragraph_format.left_indent.pt, 0.0)
        self.assertEqual(paragraph.paragraph_format.first_line_indent.pt, 0.0)

    def test_bibliography_cross_reference_is_suppressed_when_output_block_missing(
        self,
    ) -> None:
        project_root = self.create_project()
        self.update_bibliography_plan(
            project_root,
            source_mode="agent_generate_verified_only",
            output_block_present=False,
        )
        self.write_bibliography_sources(
            project_root,
            [
                {
                    "id": "ref_0001",
                    "title": "Example Paper",
                    "authors": ["Alice", "Bob"],
                    "year": "2024",
                    "container": "Journal Name",
                    "doi": "10.1000/example",
                    "evidence": {
                        "source_mode": "agent_generate_verified_only",
                        "verified_by_agent": True,
                    },
                }
            ],
        )
        (project_root / "docs" / "report_body.md").write_text(
            "[[REF:bibliography:ref_0001]]\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import docx

        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(
            item for item in redacted.paragraphs if "[[REF:bibliography:ref_0001]]" in item.text
        )

        self.assertEqual(paragraph.text.strip(), "[[REF:bibliography:ref_0001]]")
        self.assertNotIn("w:hyperlink", paragraph._p.xml)


class EquationRenderingTests(ProjectHarness):
    def test_markdown_parser_preserves_inline_equations(self) -> None:
        from scripts._report_markdown import markdown_to_blocks

        markdown_path = self.write_markdown("由 $a=b$ 可得")

        blocks = markdown_to_blocks(markdown_path)

        self.assertEqual(blocks[0]["kind"], "paragraph")
        self.assertEqual(
            blocks[0]["segments"],
            [
                {"kind": "text", "text": "由 "},
                {"kind": "inline_equation", "latex": "a=b"},
                {"kind": "text", "text": " 可得"},
            ],
        )

    def test_markdown_parser_preserves_inline_equations_in_list_items(self) -> None:
        from scripts._report_markdown import markdown_to_blocks

        markdown_path = self.write_markdown("- $s[n]$ 为原始语音样信号")

        blocks = markdown_to_blocks(markdown_path)

        self.assertEqual(blocks[0]["kind"], "list_item")
        self.assertEqual(
            blocks[0]["segments"],
            [
                {"kind": "inline_equation", "latex": "s[n]"},
                {"kind": "text", "text": " 为原始语音样信号"},
            ],
        )

    def test_markdown_parser_emits_block_equation_nodes(self) -> None:
        from scripts._report_markdown import markdown_to_blocks

        markdown_path = self.write_markdown("$$\n\\frac{a}{b}\n$$\n")

        blocks = markdown_to_blocks(markdown_path)

        self.assertEqual(blocks[0]["kind"], "equation")
        self.assertEqual(blocks[0]["latex"], "\\frac{a}{b}")

    def test_block_equation_nodes_get_stable_ids_for_reference(self) -> None:
        from scripts._report_markdown import markdown_to_blocks

        markdown_path = self.write_markdown("$$\na=b\n$$\n\n$$\nx^2\n$$\n")

        blocks = markdown_to_blocks(markdown_path)
        equations = [block for block in blocks if block["kind"] == "equation"]

        self.assertEqual([block["id"] for block in equations], ["eq_0001", "eq_0002"])

    def test_inline_equation_renders_as_word_equation_object(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "由 $a=b$ 可得。",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import docx

        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(
            item for item in redacted.paragraphs if "由" in item.text or "可得" in item.text
        )

        self.assertIn("<m:oMath", paragraph._p.xml)
        self.assertIn(">a<", paragraph._p.xml)
        self.assertIn(">b<", paragraph._p.xml)

    def test_inline_equation_supports_eta_greek_letter(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "其中 $\\eta[n]$ 表示噪声项。",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import docx

        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(item for item in redacted.paragraphs if "表示噪声项" in item.text)

        self.assertIn("<m:oMath", paragraph._p.xml)
        self.assertIn(">η<", paragraph._p.xml)
        self.assertNotIn("$\\eta[n]$", paragraph._p.xml)

    def test_list_item_inline_equation_renders_as_word_equation_object(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "### 2.1 参数说明\n\n- $s[n]$ 为原始语音样信号\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import docx

        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(item for item in redacted.paragraphs if "为原始语音样信号" in item.text)

        self.assertEqual(paragraph.style.name, "列表符号")
        self.assertIn("<m:oMath", paragraph._p.xml)
        self.assertIn(">s<", paragraph._p.xml)
        self.assertIn(">n<", paragraph._p.xml)
        self.assertNotIn("$s[n]$", paragraph._p.xml)

    def test_block_equation_renders_as_word_equation_paragraph(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "$$\n\\frac{a}{b}\n$$\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import docx

        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(item for item in redacted.paragraphs if "<m:oMath" in item._p.xml)

        self.assertIn("<m:oMathPara", paragraph._p.xml)
        self.assertIn("<m:f>", paragraph._p.xml)

    def test_block_equation_uses_parenthesized_numbering(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "$$\nx^2\n$$\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import docx

        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(item for item in redacted.paragraphs if "<m:oMath" in item._p.xml)

        self.assertIn("<m:eqArr", paragraph._p.xml)
        self.assertIn("<m:d>", paragraph._p.xml)
        self.assertIn(">#<", paragraph._p.xml)
        self.assertNotIn("<w:jc w:val=\"center\"/>", paragraph._p.xml)
        self.assertIn("bookmarkStart", paragraph._p.xml)
        self.assertIn("eq_0001", paragraph._p.xml)
        start_index = paragraph._p.xml.index("bookmarkStart")
        math_index = paragraph._p.xml.index("<m:oMath")
        end_index = paragraph._p.xml.index("bookmarkEnd")
        self.assertLess(start_index, math_index)
        self.assertLess(math_index, end_index)

    def test_equation_cross_reference_renders_formula_label(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "$$\nx^2\n$$\n\n[[REF:equation:eq_0001|由]]\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import docx

        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(item for item in redacted.paragraphs if "公式(1)" in item.text)

        self.assertEqual(paragraph.text.strip(), "由公式(1)")
        self.assertIn("w:hyperlink", paragraph._p.xml)
        self.assertIn('w:anchor="eq_0001"', paragraph._p.xml)
        self.assertIn(">公式(1)<", paragraph._p.xml)


if __name__ == "__main__":
    unittest.main()
