from __future__ import annotations

import json
import os
import shutil
import subprocess
import tempfile
import unittest
import uuid
from pathlib import Path

import docx
import yaml


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PYTHON = Path(r"D:\Miniconda\python.exe")
TEST_TEMP_ROOT = PROJECT_ROOT / "temp" / "global-default-tests"
TEST_TEMP_ROOT.mkdir(parents=True, exist_ok=True)


class RepoTemporaryDirectory:
    def __init__(self, suffix: str | None = None, prefix: str | None = None, dir: str | None = None, ignore_cleanup_errors: bool = False) -> None:
        root = Path(dir) if dir else TEST_TEMP_ROOT
        root.mkdir(parents=True, exist_ok=True)
        folder_name = f"{prefix or 'tmp'}{uuid.uuid4().hex}{suffix or ''}"
        self.path = root / folder_name
        self.path.mkdir(parents=True, exist_ok=False)
        self.name = str(self.path)

    def __enter__(self) -> str:
        return self.name

    def __exit__(self, exc_type, exc, tb) -> None:
        shutil.rmtree(self.path, ignore_errors=True)


tempfile.TemporaryDirectory = RepoTemporaryDirectory


class GlobalDefaultsTests(unittest.TestCase):
    def create_project(self) -> Path:
        project_root = TEST_TEMP_ROOT / uuid.uuid4().hex
        project_root.mkdir(parents=True, exist_ok=True)
        self.addCleanup(lambda: shutil.rmtree(project_root, ignore_errors=True))
        return project_root

    def run_workflow(self, project_root: Path, action: str, *extra_args: str, defaults_path: Path | None = None) -> subprocess.CompletedProcess[str]:
        env = os.environ.copy()
        if defaults_path is not None:
            env['GOEASY_GLOBAL_DEFAULTS_PATH'] = str(defaults_path)
        return subprocess.run(
            [
                str(PYTHON),
                str(PROJECT_ROOT / 'scripts' / 'workflow_agent.py'),
                action,
                '--project-root',
                str(project_root),
                *extra_args,
            ],
            capture_output=True,
            text=True,
            env=env,
        )

    def run_init(self, project_root: Path, *extra_args: str, defaults_path: Path | None = None) -> subprocess.CompletedProcess[str]:
        env = os.environ.copy()
        if defaults_path is not None:
            env['GOEASY_GLOBAL_DEFAULTS_PATH'] = str(defaults_path)
        return subprocess.run(
            [
                str(PYTHON),
                str(PROJECT_ROOT / 'scripts' / 'init_project.py'),
                '--project-root',
                str(project_root),
                *extra_args,
            ],
            capture_output=True,
            text=True,
            env=env,
        )

    def test_defaults_onboard_use_defaults_persists_global_defaults(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            defaults_path = project_root / 'global-defaults.json'

            result = self.run_workflow(project_root, 'defaults-onboard', '--use-defaults', defaults_path=defaults_path)

            self.assertEqual(result.returncode, 0, msg=result.stderr)
            payload = json.loads(result.stdout)
            self.assertEqual(payload['status'], 'ok')
            self.assertTrue(defaults_path.exists())
            saved = json.loads(defaults_path.read_text(encoding='utf-8'))
            self.assertIn('templates', saved)
            self.assertIn('decisions', saved)

    def test_defaults_onboard_customize_writes_defaults_preview_artifacts(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            defaults_path = project_root / 'global-defaults.json'

            result = self.run_workflow(project_root, 'defaults-onboard', '--customize', defaults_path=defaults_path)

            self.assertEqual(result.returncode, 0, msg=result.stderr)
            payload = json.loads(result.stdout)
            self.assertEqual(payload['status'], 'ok')
            self.assertTrue((project_root / 'out' / 'defaults-preview.docx').exists())
            self.assertTrue((project_root / 'out' / 'defaults-preview.summary.json').exists())
            self.assertTrue(defaults_path.exists())
            saved = json.loads(defaults_path.read_text(encoding='utf-8'))
            self.assertEqual(
                saved['templates']['template_source'],
                str((project_root / 'templates' / 'template.user.docx').resolve()),
            )
            self.assertEqual(
                saved['templates']['reference_template_source'],
                str((project_root / 'templates' / 'reference.user.docx').resolve()),
            )

    def test_init_project_seeds_template_user_from_global_defaults_source(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            defaults_path = project_root / 'global-defaults.json'
            template_source = project_root / 'custom-template.docx'
            reference_source = project_root / 'custom-reference.docx'

            doc = docx.Document()
            doc.add_paragraph('GLOBAL TEMPLATE SENTINEL')
            doc.save(template_source)

            ref_doc = docx.Document()
            ref_doc.add_paragraph('GLOBAL REFERENCE SENTINEL')
            ref_doc.save(reference_source)

            defaults_path.write_text(
                json.dumps(
                    {
                        'version': '1.0',
                        'templates': {
                            'template_source': str(template_source),
                            'reference_template_source': str(reference_source),
                        },
                        'decisions': {
                            'report_profile': 'standard',
                            'toc_enabled': False,
                            'references_required': True,
                            'appendix_enabled': False,
                            'agent_may_write_explanatory_text': True,
                            'default_template_protected': True,
                        },
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding='utf-8',
            )

            result = self.run_init(project_root, defaults_path=defaults_path)

            self.assertEqual(result.returncode, 0, msg=result.stderr)
            template_user = docx.Document(project_root / 'templates' / 'template.user.docx')
            reference_user = docx.Document(project_root / 'templates' / 'reference.user.docx')
            self.assertIn('GLOBAL TEMPLATE SENTINEL', [p.text for p in template_user.paragraphs])
            self.assertIn('GLOBAL REFERENCE SENTINEL', [p.text for p in reference_user.paragraphs])

    def test_init_project_seeds_non_null_decision_defaults_from_global_defaults(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            defaults_path = project_root / 'global-defaults.json'
            defaults_path.write_text(
                json.dumps(
                    {
                        'version': '1.0',
                        'templates': {
                            'template_source': str((PROJECT_ROOT / 'templates' / 'template.user.docx').resolve()),
                            'reference_template_source': str((PROJECT_ROOT / 'templates' / 'reference.user.docx').resolve()),
                        },
                        'decisions': {
                            'report_profile': 'body_only',
                            'toc_enabled': False,
                            'references_required': False,
                            'appendix_enabled': False,
                            'agent_may_write_explanatory_text': False,
                            'default_template_protected': False,
                        },
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding='utf-8',
            )

            result = self.run_init(project_root, defaults_path=defaults_path)

            self.assertEqual(result.returncode, 0, msg=result.stderr)
            task_contract = yaml.safe_load((project_root / 'report.task.yaml').read_text(encoding='utf-8'))
            decisions = task_contract['decisions']
            self.assertEqual(decisions['report_profile'], 'body_only')
            self.assertFalse(decisions['agent_may_write_explanatory_text'])
            self.assertFalse(decisions['default_template_protected'])

    def test_defaults_export_round_trips_imported_payload(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            defaults_path = project_root / 'global-defaults.json'
            export_path = project_root / 'exported-defaults.json'

            onboard_result = self.run_workflow(project_root, 'defaults-onboard', '--use-defaults', defaults_path=defaults_path)
            self.assertEqual(onboard_result.returncode, 0, msg=onboard_result.stderr)

            export_result = self.run_workflow(project_root, 'defaults-export', '--target', str(export_path), defaults_path=defaults_path)
            self.assertEqual(export_result.returncode, 0, msg=export_result.stderr)
            self.assertTrue(export_path.exists())
            self.assertEqual(
                json.loads(defaults_path.read_text(encoding='utf-8')),
                json.loads(export_path.read_text(encoding='utf-8')),
            )

    def test_status_reports_pair_state_and_ready_refuses_mismatched_preview_pair(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            source_template = project_root / 'voice-template.docx'
            doc = docx.Document()
            doc.add_paragraph('课程考核报告')
            doc.add_paragraph('姓 名：')
            doc.add_paragraph('学 号：')
            doc.add_paragraph('完成日期：')
            doc.add_heading('课程题目', level=1)
            doc.add_paragraph('这里是普通正文。')
            doc.save(source_template)

            init_result = self.run_init(project_root, '--template', str(source_template))
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            recommendation_path = project_root / 'logs' / 'template_style_recommendation.json'
            recommendation = json.loads(recommendation_path.read_text(encoding='utf-8'))
            recommendation['pairing']['pair_id'] = 'PAIR-A'
            recommendation_path.write_text(json.dumps(recommendation, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')

            summary_path = project_root / 'out' / 'preview.summary.json'
            summary = json.loads(summary_path.read_text(encoding='utf-8'))
            summary['pairing']['pair_id'] = 'PAIR-B'
            summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')

            status_result = self.run_workflow(project_root, 'status')
            self.assertEqual(status_result.returncode, 0, msg=status_result.stderr)
            status_payload = json.loads(status_result.stdout)
            self.assertEqual(status_payload['artifacts']['pair_state'], 'mismatched')

            ready_result = self.run_workflow(project_root, 'ready')
            self.assertEqual(ready_result.returncode, 1, msg=ready_result.stderr)
            ready_payload = json.loads(ready_result.stdout)
            issue_kinds = {item['kind'] for item in ready_payload['issues']}
            self.assertIn('mismatched_preview_pair', issue_kinds)

    def test_status_reports_stale_pair_after_runtime_authority_changes(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            source_template = project_root / 'voice-template.docx'
            doc = docx.Document()
            doc.add_paragraph('课程考核报告')
            doc.add_paragraph('姓 名：')
            doc.add_paragraph('学 号：')
            doc.add_paragraph('完成日期：')
            doc.add_heading('课程题目', level=1)
            doc.add_paragraph('这里是普通正文。')
            doc.save(source_template)

            init_result = self.run_init(project_root, '--template', str(source_template))
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            recommendation_path = project_root / 'logs' / 'template_style_recommendation.json'
            recommendation = json.loads(recommendation_path.read_text(encoding='utf-8'))
            recommendation['pending_acceptance'] = True
            recommendation_path.write_text(json.dumps(recommendation, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')

            summary_path = project_root / 'out' / 'preview.summary.json'
            summary = json.loads(summary_path.read_text(encoding='utf-8'))
            summary['template_recommendation']['pending_acceptance'] = True
            summary_path.write_text(json.dumps(summary, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')

            plan_path = project_root / 'config' / 'template.plan.json'
            plan = json.loads(plan_path.read_text(encoding='utf-8'))
            plan['selection']['primary_template'] = './templates/template.recommended.docx'
            plan_path.write_text(json.dumps(plan, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')

            status_result = self.run_workflow(project_root, 'status')
            self.assertEqual(status_result.returncode, 0, msg=status_result.stderr)
            status_payload = json.loads(status_result.stdout)
            self.assertEqual(status_payload['artifacts']['pair_state'], 'stale')

            ready_result = self.run_workflow(project_root, 'ready')
            self.assertEqual(ready_result.returncode, 1, msg=ready_result.stderr)
            ready_payload = json.loads(ready_result.stdout)
            issue_kinds = {item['kind'] for item in ready_payload['issues']}
            self.assertIn('stale_preview_pair', issue_kinds)


if __name__ == '__main__':
    unittest.main()
