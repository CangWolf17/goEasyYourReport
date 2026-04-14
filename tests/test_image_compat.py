from __future__ import annotations

import contextlib
import io
import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

import docx
from docx.text.run import Run

from scripts._image_compat import MAX_LONG_EDGE, normalize_image_for_docx
from scripts._report_render import apply_image_block


PROJECT_ROOT = Path(__file__).resolve().parents[1]


class ImageCompatTests(unittest.TestCase):
    def create_rgb_jpeg(self, path: Path, size: tuple[int, int]) -> None:
        from PIL import Image

        image = Image.new("RGB", size, (120, 180, 200))
        image.save(path, format="JPEG", quality=92)

    def create_rgba_png(self, path: Path, size: tuple[int, int]) -> None:
        from PIL import Image

        image = Image.new("RGBA", size, (120, 180, 200, 128))
        image.save(path, format="PNG")

    def test_normalize_image_for_docx_prefers_jpeg_for_photos(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            image_path = project_root / "photo.jpg"
            self.create_rgb_jpeg(image_path, (3000, 2000))

            result = normalize_image_for_docx(
                project_root,
                image_path,
                reason="UnexpectedEndOfFileError",
            )

            self.assertEqual(result.output_format, "JPEG")
            self.assertTrue(result.resized)
            self.assertLess(result.output_bytes, result.original_bytes)
            generated = Path(result.generated)
            self.assertTrue(generated.exists())
            self.assertEqual(generated.suffix.lower(), ".jpg")

            from PIL import Image

            with Image.open(generated) as normalized:
                self.assertLessEqual(max(normalized.size), MAX_LONG_EDGE)

    def test_normalize_image_for_docx_preserves_png_for_alpha(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            image_path = project_root / "alpha.png"
            self.create_rgba_png(image_path, (800, 600))

            result = normalize_image_for_docx(
                project_root,
                image_path,
                reason="UnexpectedEndOfFileError",
            )

            self.assertEqual(result.output_format, "PNG")
            self.assertFalse(result.resized)
            self.assertTrue(Path(result.generated).suffix.lower() == ".png")

    def test_apply_image_block_falls_back_to_generated_asset(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            docs_dir = project_root / "docs"
            docs_dir.mkdir(parents=True, exist_ok=True)
            image_path = docs_dir / "photo.jpg"
            self.create_rgb_jpeg(image_path, (2600, 1800))

            document = docx.Document()
            paragraph = document.add_paragraph()
            image_status = {"inserted": [], "normalized": [], "failed": []}
            block = {"alt": "Photo", "path": "photo.jpg"}

            original_add_picture = Run.add_picture

            def flaky_add_picture(run_self, image_file, *args, **kwargs):
                if str(image_file) == str(image_path):
                    from docx.image.exceptions import UnexpectedEndOfFileError

                    raise UnexpectedEndOfFileError()
                return original_add_picture(run_self, image_file, *args, **kwargs)

            with mock.patch.object(Run, "add_picture", new=flaky_add_picture):
                image_paragraph, inserted = apply_image_block(
                    paragraph,
                    block,
                    {"Caption"},
                    None,
                    docs_dir,
                    image_status,
                    project_root,
                )

            self.assertTrue(inserted)
            self.assertIs(image_paragraph, paragraph)
            self.assertEqual(image_status["failed"], [])
            self.assertEqual(len(image_status["normalized"]), 1)
            normalized = image_status["normalized"][0]
            self.assertEqual(normalized["reason"], "UnexpectedEndOfFileError")
            self.assertEqual(normalized["output_format"], "JPEG")
            self.assertTrue(Path(normalized["generated_path"]).exists())

    def test_apply_image_block_missing_file_still_fails_normally(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            docs_dir = project_root / "docs"
            docs_dir.mkdir(parents=True, exist_ok=True)

            document = docx.Document()
            paragraph = document.add_paragraph()
            image_status = {"inserted": [], "normalized": [], "failed": []}
            block = {"alt": "Missing", "path": "missing.jpg"}

            _paragraph, inserted = apply_image_block(
                paragraph,
                block,
                {"Caption"},
                None,
                docs_dir,
                image_status,
                project_root,
            )

            self.assertFalse(inserted)
            self.assertEqual(image_status["normalized"], [])
            self.assertEqual(len(image_status["failed"]), 1)

    def test_build_report_payload_surfaces_normalized_images(self) -> None:
        from scripts import build_report

        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    sys.executable,
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            image_path = project_root / "docs" / "images" / "photo.jpg"
            image_path.parent.mkdir(parents=True, exist_ok=True)
            self.create_rgb_jpeg(image_path, (2600, 1800))
            (project_root / "docs" / "report_body.md").write_text(
                "## Figures\n\n![Photo](images/photo.jpg)\n",
                encoding="utf-8",
            )

            original_add_picture = Run.add_picture

            def flaky_add_picture(run_self, image_file, *args, **kwargs):
                if str(image_file) == str(image_path):
                    from docx.image.exceptions import UnexpectedEndOfFileError

                    raise UnexpectedEndOfFileError()
                return original_add_picture(run_self, image_file, *args, **kwargs)

            with mock.patch.object(Run, "add_picture", new=flaky_add_picture):
                stdout = io.StringIO()
                with mock.patch.object(
                    sys,
                    "argv",
                    ["build_report.py", "--project-root", str(project_root)],
                ), contextlib.redirect_stdout(stdout):
                    exit_code = build_report.main()

            self.assertEqual(exit_code, 0)
            payload = json.loads(stdout.getvalue())
            self.assertEqual(payload["images"]["failed"], [])
            self.assertEqual(len(payload["images"]["normalized"]), 1)
            normalized = payload["images"]["normalized"][0]
            self.assertEqual(normalized["reason"], "UnexpectedEndOfFileError")
            self.assertEqual(normalized["output_format"], "JPEG")
            self.assertTrue(Path(normalized["generated_path"]).exists())


if __name__ == "__main__":
    unittest.main()
