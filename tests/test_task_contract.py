from __future__ import annotations

import json
import subprocess
import unittest
import uuid
from pathlib import Path

import docx
import yaml


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PYTHON = Path(r"D:\Miniconda\python.exe")


class TaskContractTests(unittest.TestCase):
    def create_project(self) -> Path:
        sandbox_root = PROJECT_ROOT / "temp" / "task-contract-tests"
        sandbox_root.mkdir(parents=True, exist_ok=True)
        project_root = sandbox_root / uuid.uuid4().hex
        project_root.mkdir(parents=True, exist_ok=True)
        self.addCleanup(
            lambda: __import__("shutil").rmtree(project_root, ignore_errors=True)
        )
        return project_root

    def load_task_yaml(self, project_root: Path) -> dict[str, object]:
        return yaml.safe_load(
            (project_root / "report.task.yaml").read_text(encoding="utf-8")
        )

    def dump_task_yaml(self, project_root: Path, payload: dict[str, object]) -> None:
        (project_root / "report.task.yaml").write_text(
            yaml.safe_dump(payload, allow_unicode=True, sort_keys=False),
            encoding="utf-8",
        )

    def make_current_ready_to_build(self, project_root: Path) -> None:
        plan_path = project_root / "config" / "template.plan.json"
        plan = json.loads(plan_path.read_text(encoding="utf-8"))
        plan["semantics"]["cross_references"]["figure_table_enabled"] = False
        plan["semantics"]["bibliography"]["source_mode"] = "user_supplied_files"
        plan_path.write_text(
            json.dumps(plan, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        template_path = project_root / "templates" / "template.user.docx"
        template = docx.Document(template_path)
        for paragraph in list(template.paragraphs[0:4]):
            paragraph._element.getparent().remove(paragraph._element)
        template.save(template_path)

        prepare_result = self.run_workflow(project_root, "prepare")
        self.assertEqual(prepare_result.returncode, 0, msg=prepare_result.stderr)
        ready_result = self.run_workflow(project_root, "ready")
        self.assertEqual(ready_result.returncode, 0, msg=ready_result.stderr)

    def insert_toc_placeholder(self, project_root: Path) -> None:
        template_path = project_root / "templates" / "template.user.docx"
        template = docx.Document(template_path)
        anchor = next(
            paragraph
            for paragraph in template.paragraphs
            if getattr(paragraph.style, "name", "") == "标题2"
        )
        anchor.insert_paragraph_before("目录")
        template.save(template_path)

    def load_preview_summary(self, project_root: Path) -> dict[str, object]:
        return json.loads(
            (project_root / "out" / "preview.summary.json").read_text(
                encoding="utf-8"
            )
        )

    def init_project(self, project_root: Path) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [
                str(PYTHON),
                str(PROJECT_ROOT / "scripts" / "init_project.py"),
                "--project-root",
                str(project_root),
            ],
            capture_output=True,
            text=True,
        )

    def run_workflow(
        self, project_root: Path, action: str, *extra_args: str
    ) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [
                str(PYTHON),
                str(PROJECT_ROOT / "scripts" / "workflow_agent.py"),
                action,
                "--project-root",
                str(project_root),
                *extra_args,
            ],
            capture_output=True,
            text=True,
        )

    def assert_task_contract_shape(self, payload: dict[str, object]) -> None:
        for key in (
            "schema",
            "task",
            "requirements",
            "inputs",
            "decisions",
            "runtime",
        ):
            self.assertIn(key, payload)
        self.assertEqual(payload["schema"]["kind"], "report_task")
        self.assertFalse(payload["task"]["ready_to_write"])
        self.assertEqual(payload["runtime"]["workflow_config"], "./workflow.json")
        self.assertEqual(payload["runtime"]["preview_review_status"], "unknown")
        self.assertEqual(payload["runtime"]["redacted_verify_status"], "unknown")
        self.assertEqual(payload["runtime"]["acceptance_status"], "unknown")

    def test_init_project_creates_report_task_yaml(self) -> None:
        project_root = self.create_project()

        result = self.init_project(project_root)

        self.assertEqual(result.returncode, 0, msg=result.stderr)
        self.assertTrue((project_root / "report.task.yaml").exists())
        self.assert_task_contract_shape(self.load_task_yaml(project_root))

    def test_load_task_contract_returns_default_shape(self) -> None:
        from scripts._task_contract import load_task_contract

        project_root = self.create_project()
        task_path = project_root / "report.task.yaml"

        payload = load_task_contract(task_path)

        self.assert_task_contract_shape(payload)

    def test_dump_task_contract_round_trips_yaml(self) -> None:
        from scripts._task_contract import (
            default_task_contract,
            dump_task_contract,
            load_task_contract,
        )

        project_root = self.create_project()
        task_path = project_root / "report.task.yaml"
        payload = default_task_contract()
        payload["task"]["name"] = "示例任务"

        dump_task_contract(task_path, payload)
        loaded = load_task_contract(task_path)

        self.assert_task_contract_shape(loaded)
        self.assertEqual(loaded["task"]["name"], "示例任务")

    def test_sync_prepare_task_contract_does_not_announce_build_before_ready(self) -> None:
        from scripts._task_contract import default_task_contract, dump_task_contract
        from scripts.workflow_agent import sync_prepare_task_contract

        project_root = self.create_project()
        dump_task_contract(project_root / "report.task.yaml", default_task_contract())

        task_contract = sync_prepare_task_contract(project_root, [])

        self.assertFalse(task_contract["task"]["ready_to_write"])
        self.assertTrue(task_contract["task"]["needs_user_input"])
        self.assertEqual(task_contract["runtime"]["next_step"], "resolve_report_task_gate")

    def test_workflow_agent_prepare_updates_report_task_runtime(self) -> None:
        project_root = self.create_project()

        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        result = self.run_workflow(project_root, "prepare")

        self.assertEqual(result.returncode, 0, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["action"], "prepare")

        task_contract = self.load_task_yaml(project_root)
        runtime = task_contract["runtime"]
        self.assertEqual(runtime["preview_output"], "./out/preview.docx")
        self.assertEqual(runtime["semantic_preview_output"], "./out/semantic-preview.docx")
        self.assertEqual(runtime["template_plan"], "./config/template.plan.json")
        self.assertEqual(runtime["field_binding"], "./config/field.binding.json")
        self.assertEqual(payload["next_step"], runtime["next_step"])
        self.assertIn(runtime["preview_review_status"], {"pass", "needs_user_decision", "needs_preview_revision"})

        summary = self.load_preview_summary(project_root)
        self.assertEqual(summary["task_contract"]["next_step"], runtime["next_step"])
        self.assertEqual(summary["task_contract"]["stage"], task_contract["task"]["stage"])

    def test_workflow_agent_build_blocks_when_report_task_not_ready(self) -> None:
        project_root = self.create_project()

        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        result = self.run_workflow(project_root, "build")

        self.assertEqual(result.returncode, 1, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assertIn(
            payload["status"],
            {"needs_user_confirmation", "needs_agent_handoff"},
        )
        self.assertTrue(
            "ready_to_write" in payload["summary"]
            or any(
                "ready_to_write" in str(item)
                for item in payload.get("issues", [])
            )
        )
        self.assertFalse((project_root / "out" / "redacted.docx").exists())

    def test_workflow_agent_build_allows_current_ready_to_write_task(self) -> None:
        project_root = self.create_project()

        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)
        self.make_current_ready_to_build(project_root)

        result = self.run_workflow(project_root, "build")

        self.assertEqual(result.returncode, 0, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["status"], "ok")
        updated = self.load_task_yaml(project_root)
        self.assertEqual(updated["task"]["stage"], "redacted_built")
        self.assertEqual(updated["runtime"]["redacted_output"], "./out/redacted.docx")
        self.assertEqual(updated["runtime"]["next_step"], "verify")

    def test_workflow_agent_ready_requires_preview_summary_first(self) -> None:
        project_root = self.create_project()

        result = self.run_workflow(project_root, "ready")

        self.assertEqual(result.returncode, 1, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["action"], "ready")
        self.assertEqual(payload["status"], "needs_user_confirmation")
        self.assertEqual(payload["issues"][0]["kind"], "missing_preview_summary")

    def test_workflow_agent_ready_marks_task_ready_after_clean_prepare(self) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        plan_path = project_root / "config" / "template.plan.json"
        plan = json.loads(plan_path.read_text(encoding="utf-8"))
        plan["semantics"]["cross_references"]["figure_table_enabled"] = False
        plan["semantics"]["bibliography"]["source_mode"] = "user_supplied_files"
        plan_path.write_text(
            json.dumps(plan, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        template_path = project_root / "templates" / "template.user.docx"
        template = docx.Document(template_path)
        for paragraph in list(template.paragraphs[0:4]):
            paragraph._element.getparent().remove(paragraph._element)
        template.save(template_path)

        prepare_result = self.run_workflow(project_root, "prepare")
        self.assertEqual(prepare_result.returncode, 0, msg=prepare_result.stderr)

        ready_result = self.run_workflow(project_root, "ready")
        self.assertEqual(ready_result.returncode, 0, msg=ready_result.stderr)
        payload = json.loads(ready_result.stdout)
        self.assertEqual(payload["action"], "ready")
        self.assertEqual(payload["status"], "ok")

        task_contract = self.load_task_yaml(project_root)
        self.assertTrue(task_contract["task"]["ready_to_write"])
        self.assertEqual(task_contract["runtime"]["next_step"], "build")

    def test_workflow_agent_build_invalidates_ready_when_preview_inputs_change(self) -> None:
        project_root = self.create_project()
        self.assertEqual(self.init_project(project_root).returncode, 0)
        self.make_current_ready_to_build(project_root)

        (project_root / "docs" / "report_body.md").write_text(
            "# Updated Body\n\n## New Section\n\nBody changed after ready.\n",
            encoding="utf-8",
        )

        result = self.run_workflow(project_root, "build")

        self.assertEqual(result.returncode, 1, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["action"], "build")
        self.assertEqual(payload["status"], "needs_agent_handoff")
        self.assertEqual(payload["next_step"], "preview")
        self.assertTrue(
            any(issue["kind"] == "stale_semantic_preview" for issue in payload["issues"])
        )

        task_contract = self.load_task_yaml(project_root)
        self.assertFalse(task_contract["task"]["ready_to_write"])
        self.assertEqual(task_contract["runtime"]["next_step"], "preview")

    def test_workflow_agent_build_invalidates_generated_semantic_preview_when_requirements_change(self) -> None:
        project_root = self.create_project()
        self.assertEqual(self.init_project(project_root).returncode, 0)
        (project_root / "docs" / "report_body.md").write_text(
            "# 摘要\n",
            encoding="utf-8",
        )
        self.make_current_ready_to_build(project_root)
        summary = self.load_preview_summary(project_root)
        self.assertEqual(summary["semantic_preview"]["scaffold_mode"], "generated")

        (project_root / "docs" / "task_requirements.md").write_text(
            "- 新增评分点：必须强调实验结论。\n",
            encoding="utf-8",
        )

        result = self.run_workflow(project_root, "build")

        self.assertEqual(result.returncode, 1, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["status"], "needs_agent_handoff")
        self.assertTrue(
            any(issue["kind"] == "stale_semantic_preview" for issue in payload["issues"])
        )

    def test_workflow_agent_ready_refuses_when_blocking_confirmations_remain(self) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        self.insert_toc_placeholder(project_root)
        prepare_result = self.run_workflow(project_root, "prepare")
        self.assertEqual(prepare_result.returncode, 0, msg=prepare_result.stderr)

        ready_result = self.run_workflow(project_root, "ready")
        self.assertEqual(ready_result.returncode, 1, msg=ready_result.stderr)
        payload = json.loads(ready_result.stdout)
        self.assertEqual(payload["status"], "needs_user_confirmation")
        self.assertTrue(
            any(issue["kind"] == "confirmation_required" for issue in payload["issues"])
        )

    def test_workflow_agent_status_reports_missing_preview_summary(self) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)
        (project_root / "out" / "preview.summary.json").unlink()

        result = self.run_workflow(project_root, "status")

        self.assertEqual(result.returncode, 0, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["action"], "status")
        self.assertEqual(payload["status"], "ok")
        self.assertIn("no preview summary yet", payload["summary"])
        self.assertIn(
            "run prepare to generate preview.summary and confirmation details",
            payload["warnings"],
        )

    def test_workflow_agent_status_reports_blocking_and_advisory_items(self) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        prepare_result = self.run_workflow(project_root, "prepare")
        self.assertEqual(prepare_result.returncode, 0, msg=prepare_result.stderr)

        result = self.run_workflow(project_root, "status")

        self.assertEqual(result.returncode, 0, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["action"], "status")
        self.assertEqual(payload["status"], "ok")
        self.assertTrue(
            any(issue["kind"] == "decision_required" for issue in payload["issues"])
        )

    def test_task_contract_seeds_high_level_decision_defaults(self) -> None:
        project_root = self.create_project()

        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        decisions = self.load_task_yaml(project_root)["decisions"]
        self.assertEqual(decisions["report_profile"], "standard")
        self.assertIsNone(decisions["toc_enabled"])
        self.assertIsNone(decisions["references_required"])
        self.assertIsNone(decisions["appendix_enabled"])
        self.assertTrue(decisions["agent_may_write_explanatory_text"])
        self.assertTrue(decisions["default_template_protected"])

    def test_preview_summary_includes_high_level_report_decisions(self) -> None:
        project_root = self.create_project()

        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        summary = self.load_preview_summary(project_root)
        self.assertIn("report_decisions", summary)
        self.assertIsNone(summary["report_decisions"]["toc_enabled"])
        self.assertTrue(
            summary["report_decisions"]["default_template_protected"]
        )

    def test_workflow_agent_preview_updates_task_contract_runtime_and_summary(self) -> None:
        project_root = self.create_project()

        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        result = self.run_workflow(project_root, "preview")

        self.assertEqual(result.returncode, 1, msg=result.stderr)
        payload = json.loads(result.stdout)
        task_contract = self.load_task_yaml(project_root)
        summary = self.load_preview_summary(project_root)

        self.assertEqual(payload["next_step"], task_contract["runtime"]["next_step"])
        self.assertEqual(summary["task_contract"]["next_step"], payload["next_step"])
        self.assertEqual(summary["task_contract"]["stage"], task_contract["task"]["stage"])
        self.assertEqual(
            summary["task_contract"]["ready_to_write"],
            task_contract["task"]["ready_to_write"],
        )
        self.assertTrue((project_root / "out" / "semantic-preview.docx").exists())
        self.assertIn("semantic_preview", summary)
        self.assertIn(
            task_contract["runtime"]["preview_review_status"],
            {"pass", "needs_user_decision", "needs_preview_revision"},
        )
        if task_contract["runtime"]["preview_review_status"] != "pass":
            self.assertEqual(payload["next_step"], "review_preview_summary")

    def test_workflow_agent_bootstrap_initializes_and_prepares_external_target(self) -> None:
        project_root = self.create_project()

        result = self.run_workflow(project_root, "bootstrap")

        self.assertEqual(result.returncode, 0, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["action"], "bootstrap")
        self.assertTrue((project_root / "workflow.json").exists())
        self.assertTrue((project_root / "report.task.yaml").exists())
        self.assertIn("preview_summary", payload["artifacts"])

    def test_prepare_syncs_template_mirrors_from_plan(self) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        recommended_template = project_root / "templates" / "template.recommended.docx"
        recommended_template.write_bytes(
            (project_root / "templates" / "template.user.docx").read_bytes()
        )

        plan_path = project_root / "config" / "template.plan.json"
        plan = json.loads(plan_path.read_text(encoding="utf-8"))
        plan["selection"]["primary_template"] = "./templates/template.recommended.docx"
        plan_path.write_text(
            json.dumps(plan, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

        workflow_path = project_root / "workflow.json"
        workflow = json.loads(workflow_path.read_text(encoding="utf-8"))
        workflow["templates"]["main_template"] = "./templates/template.user.docx"
        workflow_path.write_text(
            json.dumps(workflow, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

        task_contract = self.load_task_yaml(project_root)
        task_contract["inputs"]["template_path"] = "./templates/template.user.docx"
        self.dump_task_yaml(project_root, task_contract)

        result = self.run_workflow(project_root, "prepare")

        self.assertEqual(result.returncode, 0, msg=result.stderr)
        workflow = json.loads(workflow_path.read_text(encoding="utf-8"))
        self.assertEqual(
            workflow["templates"]["main_template"],
            "./templates/template.recommended.docx",
        )
        task_contract = self.load_task_yaml(project_root)
        self.assertEqual(
            task_contract["inputs"]["template_path"],
            "./templates/template.recommended.docx",
        )

    def test_build_preview_marks_missing_field_candidates_as_advisory_for_body_only_template(
        self,
    ) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        template_path = project_root / "templates" / "template.user.docx"
        template = docx.Document(template_path)
        for paragraph in list(template.paragraphs[0:4]):
            paragraph._element.getparent().remove(paragraph._element)
        template.save(template_path)

        result = self.run_workflow(project_root, "prepare")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        summary = self.load_preview_summary(project_root)
        self.assertIn("no field candidates detected", summary["review"]["warnings"])
        self.assertNotIn(
            "no field candidates detected", summary["review"]["needs_confirmation"]
        )

    def test_build_preview_keeps_missing_field_candidates_blocking_when_cover_region_exists(
        self,
    ) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        template_path = project_root / "templates" / "template.user.docx"
        template = docx.Document(template_path)
        for paragraph in list(template.paragraphs[1:4]):
            paragraph._element.getparent().remove(paragraph._element)
        template.save(template_path)

        result = self.run_workflow(project_root, "prepare")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        summary = self.load_preview_summary(project_root)
        self.assertIn(
            "cover region detected without recognizable field candidates",
            summary["review"]["needs_confirmation"],
        )

    def test_body_only_profile_downgrades_cover_field_noise_to_advisory(self) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        task = self.load_task_yaml(project_root)
        task["decisions"]["report_profile"] = "body_only"
        self.dump_task_yaml(project_root, task)

        template_path = project_root / "templates" / "template.user.docx"
        template = docx.Document(template_path)
        for paragraph in list(template.paragraphs[1:4]):
            paragraph._element.getparent().remove(paragraph._element)
        template.save(template_path)

        result = self.run_workflow(project_root, "prepare")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        summary = self.load_preview_summary(project_root)
        self.assertIn("no field candidates detected", summary["review"]["warnings"])
        self.assertNotIn(
            "cover region detected without recognizable field candidates",
            summary["review"]["needs_confirmation"],
        )
        payload = json.loads(result.stdout)
        self.assertIn("body_only profile", payload["summary"])

    def test_verify_redacted_routes_to_review_and_persists_fingerprint(self) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)
        self.make_current_ready_to_build(project_root)

        build_result = self.run_workflow(project_root, "build")
        self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

        verify_result = self.run_workflow(project_root, "verify", "--target", "redacted")
        self.assertEqual(verify_result.returncode, 0, msg=verify_result.stderr)
        payload = json.loads(verify_result.stdout)
        self.assertEqual(payload["next_step"], "review")

        task_contract = self.load_task_yaml(project_root)
        self.assertEqual(task_contract["task"]["stage"], "awaiting_acceptance_review")
        self.assertEqual(task_contract["runtime"]["next_step"], "review")
        self.assertEqual(task_contract["runtime"]["redacted_verify_status"], "pass")
        self.assertTrue(task_contract["runtime"]["redacted_verify_fingerprint"])
        self.assertTrue((project_root / "out" / "_internal" / "redacted-verify.json").exists())

    def test_review_refuses_without_current_verified_redacted_output(self) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        review_result = self.run_workflow(project_root, "review")
        self.assertEqual(review_result.returncode, 1, msg=review_result.stderr)
        payload = json.loads(review_result.stdout)
        self.assertEqual(payload["action"], "review")
        self.assertEqual(payload["status"], "needs_agent_handoff")
        self.assertTrue(any(issue["kind"] == "missing_redacted_verify_pass" for issue in payload["issues"]))

    def test_review_uses_stubbed_worker_output_and_unlocks_inject(self) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)
        self.make_current_ready_to_build(project_root)

        build_result = self.run_workflow(project_root, "build")
        self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)
        verify_result = self.run_workflow(project_root, "verify", "--target", "redacted")
        self.assertEqual(verify_result.returncode, 0, msg=verify_result.stderr)

        redacted_fingerprint = self.load_task_yaml(project_root)["runtime"]["redacted_verify_fingerprint"]
        internal_root = project_root / "out" / "_internal"
        internal_root.mkdir(parents=True, exist_ok=True)
        (internal_root / "review-worker-output.json").write_text(
            json.dumps(
                {
                    "status": "pass",
                    "target_fingerprint": redacted_fingerprint,
                    "requirements_alignment": "pass",
                    "style_alignment": "pass",
                    "document_quality": "pass",
                    "preview_consistency": "pass",
                    "rerender_target": "none",
                    "blocking_findings": [],
                    "needs_decision": [],
                    "evidence": ["stubbed review output"],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        review_result = self.run_workflow(project_root, "review")
        self.assertEqual(review_result.returncode, 0, msg=review_result.stderr)
        payload = json.loads(review_result.stdout)
        self.assertEqual(payload["status"], "ok")

        task_contract = self.load_task_yaml(project_root)
        self.assertEqual(task_contract["runtime"]["acceptance_status"], "pass")
        self.assertEqual(task_contract["runtime"]["next_step"], "inject")
        self.assertEqual(
            task_contract["runtime"]["accepted_redacted_fingerprint"],
            redacted_fingerprint,
        )
        self.assertTrue((project_root / "out" / "acceptance-review.json").exists())

    def test_review_rejects_worker_output_without_target_fingerprint(self) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)
        self.make_current_ready_to_build(project_root)

        self.assertEqual(self.run_workflow(project_root, "build").returncode, 0)
        self.assertEqual(
            self.run_workflow(project_root, "verify", "--target", "redacted").returncode,
            0,
        )

        internal_root = project_root / "out" / "_internal"
        internal_root.mkdir(parents=True, exist_ok=True)
        (internal_root / "review-worker-output.json").write_text(
            json.dumps({"status": "pass"}, ensure_ascii=False),
            encoding="utf-8",
        )

        review_result = self.run_workflow(project_root, "review")
        self.assertEqual(review_result.returncode, 2, msg=review_result.stderr)
        payload = json.loads(review_result.stdout)
        self.assertEqual(payload["status"], "error")

    def test_inject_blocks_without_current_acceptance_pass(self) -> None:
        project_root = self.create_project()
        init_result = self.init_project(project_root)
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        result = self.run_workflow(project_root, "inject")
        self.assertEqual(result.returncode, 1, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["action"], "inject")
        self.assertEqual(payload["status"], "needs_agent_handoff")
        self.assertTrue(any(issue["kind"] == "acceptance_not_passed" for issue in payload["issues"]))


if __name__ == "__main__":
    unittest.main()
