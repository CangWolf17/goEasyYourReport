from __future__ import annotations

import base64
import json
import re
import shutil
import subprocess
import tempfile
import unittest
import uuid
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

import docx
import yaml
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PYTHON = Path(r"D:\Miniconda\python.exe")
TEST_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
)
TEST_TEMP_ROOT = PROJECT_ROOT / "temp" / "init-project-tests"
TEST_TEMP_ROOT.mkdir(parents=True, exist_ok=True)
STYLE_XML_DECLARATION_PATTERN = re.compile(rb"^<\?xml[^?]*\?>")
STYLE_XML_ROOT_PATTERN = re.compile(rb"<(?:\w+:)?styles\b[^>]*>")
STYLE_XML_NAMESPACES = {
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
    "w16": "http://schemas.microsoft.com/office/word/2018/wordml",
    "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
    "w16cid": "http://schemas.microsoft.com/office/word/2016/wordml/cid",
    "w16sdtdh": "http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash",
    "w16se": "http://schemas.microsoft.com/office/word/2015/wordml/symex",
}
WORD_XML_NAMESPACES = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


class RepoTemporaryDirectory:
    def __init__(
        self,
        suffix: str | None = None,
        prefix: str | None = None,
        dir: str | None = None,
        ignore_cleanup_errors: bool = False,
    ) -> None:
        root = Path(dir) if dir else TEST_TEMP_ROOT
        root.mkdir(parents=True, exist_ok=True)
        self._ignore_cleanup_errors = ignore_cleanup_errors
        folder_name = f"{prefix or 'tmp'}{uuid.uuid4().hex}{suffix or ''}"
        self.path = root / folder_name
        self.path.mkdir(parents=True, exist_ok=False)
        self.name = str(self.path)

    def __enter__(self) -> str:
        return self.name

    def cleanup(self) -> None:
        shutil.rmtree(self.path, ignore_errors=True)

    def __exit__(self, exc_type, exc, tb) -> None:
        self.cleanup()


tempfile.TemporaryDirectory = RepoTemporaryDirectory


def write_style_poor_template(path: Path) -> None:
    template_doc = docx.Document()
    template_doc.add_paragraph("课程考核报告")
    template_doc.add_paragraph("姓 名：")
    template_doc.add_paragraph("学 号：")
    template_doc.add_paragraph("完成日期：")
    template_doc.add_heading("课程题目", level=1)
    template_doc.add_heading("1 实验目的", level=2)
    template_doc.add_paragraph("这里是普通正文。")
    template_doc.add_heading("2 实验结果", level=2)
    template_doc.add_paragraph("这里是普通正文。")
    template_doc.save(path)


def strip_list_styles(path: Path) -> None:
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def qn_local(local_name: str) -> str:
        return f"{{{w_ns}}}{local_name}"

    target_names = {
        "List Bullet",
        "List Bullet 2",
        "List Bullet 3",
        "List Number",
        "List Number 2",
        "List Number 3",
        "List Continue",
        "List Continue 2",
        "List Continue 3",
        "List Paragraph",
    }

    with zipfile.ZipFile(path, "r") as source_zip:
        entries = {
            info.filename: source_zip.read(info.filename)
            for info in source_zip.infolist()
        }

    original_styles_xml = entries["word/styles.xml"]
    styles_root = ET.fromstring(original_styles_xml)
    for child in list(styles_root):
        if child.tag == qn_local("style"):
            name_element = child.find(qn_local("name"))
            style_name = (
                name_element.get(qn_local("val")) if name_element is not None else None
            )
            if style_name in target_names:
                styles_root.remove(child)
        elif child.tag == qn_local("latentStyles"):
            for latent in list(child):
                if (
                    latent.tag == qn_local("lsdException")
                    and latent.get(qn_local("name")) in target_names
                ):
                    child.remove(latent)

    for prefix, uri in STYLE_XML_NAMESPACES.items():
        ET.register_namespace(prefix, uri)

    serialized = ET.tostring(styles_root, encoding="utf-8", xml_declaration=True)
    original_declaration = STYLE_XML_DECLARATION_PATTERN.search(original_styles_xml)
    if original_declaration is not None:
        serialized = STYLE_XML_DECLARATION_PATTERN.sub(
            original_declaration.group(0),
            serialized,
            count=1,
        )
    original_root = STYLE_XML_ROOT_PATTERN.search(original_styles_xml)
    serialized_root = STYLE_XML_ROOT_PATTERN.search(serialized)
    if original_root is not None and serialized_root is not None:
        serialized = (
            serialized[: serialized_root.start()]
            + original_root.group(0)
            + serialized[serialized_root.end() :]
        )

    entries["word/styles.xml"] = serialized
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as output_zip:
        for filename, content in entries.items():
            output_zip.writestr(filename, content)


def style_font_settings(style) -> dict[str, object]:
    rpr = style.element.find(qn("w:rPr"))
    rfonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
    size = rpr.find(qn("w:sz")) if rpr is not None else None
    return {
        "ascii": None if rfonts is None else rfonts.get(qn("w:ascii")),
        "hAnsi": None if rfonts is None else rfonts.get(qn("w:hAnsi")),
        "eastAsia": None if rfonts is None else rfonts.get(qn("w:eastAsia")),
        "size": None if size is None else size.get(qn("w:val")),
    }


def run_font_settings(run) -> dict[str, object]:
    rpr = run._r.find(qn("w:rPr"))
    rfonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
    size = rpr.find(qn("w:sz")) if rpr is not None else None
    return {
        "ascii": None if rfonts is None else rfonts.get(qn("w:ascii")),
        "hAnsi": None if rfonts is None else rfonts.get(qn("w:hAnsi")),
        "eastAsia": None if rfonts is None else rfonts.get(qn("w:eastAsia")),
        "size": None if size is None else size.get(qn("w:val")),
    }


def paragraph_numbering_id(paragraph) -> str | None:
    num_id = paragraph._p.find(".//w:numPr/w:numId", WORD_XML_NAMESPACES)
    if num_id is None:
        return None
    return num_id.get(qn("w:val"))


def numbering_start_overrides(path: Path) -> dict[str, int]:
    with zipfile.ZipFile(path, "r") as docx_zip:
        numbering_xml = docx_zip.read("word/numbering.xml")

    root = ET.fromstring(numbering_xml)
    overrides: dict[str, int] = {}
    for num in root.findall("w:num", WORD_XML_NAMESPACES):
        num_id = num.get(qn("w:numId"))
        start_override = num.find(
            "w:lvlOverride/w:startOverride", WORD_XML_NAMESPACES
        )
        if num_id and start_override is not None:
            value = start_override.get(qn("w:val"))
            if value is not None:
                overrides[num_id] = int(value)
    return overrides


def set_explicit_run_font(
    run,
    *,
    ascii_name: str,
    east_asia_name: str | None,
    size_pt: float,
) -> None:
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    run.font.name = ascii_name
    run.font.size = Pt(size_pt)

    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_name)
    rfonts.set(qn("w:hAnsi"), ascii_name)
    if east_asia_name:
        rfonts.set(qn("w:eastAsia"), east_asia_name)

    size_value = str(int(size_pt * 2))
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rpr.append(sz)
    sz.set(qn("w:val"), size_value)

    sz_cs = rpr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        rpr.append(sz_cs)
    sz_cs.set(qn("w:val"), size_value)


def write_styled_cover_template(path: Path) -> None:
    template_doc = docx.Document()
    title = template_doc.add_paragraph()
    title_run = title.add_run("课程考核报告")
    title_run.bold = True
    set_explicit_run_font(
        title_run,
        ascii_name="Times New Roman",
        east_asia_name="宋体",
        size_pt=18.0,
    )

    name_paragraph = template_doc.add_paragraph()
    name_label = name_paragraph.add_run("学 生 姓 名：")
    set_explicit_run_font(
        name_label,
        ascii_name="Times New Roman",
        east_asia_name="Times New Roman",
        size_pt=15.0,
    )
    name_value = name_paragraph.add_run("XXX")
    set_explicit_run_font(
        name_value,
        ascii_name="Times New Roman",
        east_asia_name="仿宋",
        size_pt=14.0,
    )

    student_id_paragraph = template_doc.add_paragraph()
    student_id_label = student_id_paragraph.add_run("学 号：")
    set_explicit_run_font(
        student_id_label,
        ascii_name="Times New Roman",
        east_asia_name="Times New Roman",
        size_pt=15.0,
    )
    student_id_value = student_id_paragraph.add_run("XXX")
    set_explicit_run_font(
        student_id_value,
        ascii_name="Consolas",
        east_asia_name="Times New Roman",
        size_pt=13.0,
    )

    completion_paragraph = template_doc.add_paragraph()
    completion_label = completion_paragraph.add_run("完成日期：")
    set_explicit_run_font(
        completion_label,
        ascii_name="Times New Roman",
        east_asia_name="宋体",
        size_pt=15.0,
    )
    completion_value = completion_paragraph.add_run("202X年XX月XX日")
    set_explicit_run_font(
        completion_value,
        ascii_name="Times New Roman",
        east_asia_name="楷体",
        size_pt=14.0,
    )

    template_doc.add_heading("正文开始", level=1)
    template_doc.add_paragraph("这里是普通正文。")
    template_doc.save(path)


def write_cover_template_with_unbound_candidates(path: Path) -> None:
    template_doc = docx.Document()
    for line in (
        "学 院：",
        "专 业：",
        "学 生 姓 名：",
        "学 号：",
        "完成日期：",
    ):
        template_doc.add_paragraph(line)
    template_doc.add_heading("正文开始", level=1)
    template_doc.add_paragraph("这里是普通正文。")
    template_doc.save(path)


def write_cover_template_with_placeholder_candidates(path: Path) -> None:
    template_doc = docx.Document()
    for line in (
        "学       院：XXXX学院",
        "专       业：XXXX",
        "学 生 姓 名：XXX",
        "学       号：XXX",
        "评 阅 教 师：XXX",
        "完 成 时 间：202X年XX月XX日",
    ):
        template_doc.add_paragraph(line)
    template_doc.add_heading("正文开始", level=1)
    template_doc.add_paragraph("这里是普通正文。")
    template_doc.save(path)


class InitProjectTests(unittest.TestCase):
    def test_run_optional_reads_utf8_child_output(self) -> None:
        from scripts.init_project import run_optional

        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            scripts_root = project_root / "scripts"
            scripts_root.mkdir(parents=True, exist_ok=True)
            script_path = scripts_root / "emit_utf8.py"
            script_path.write_text(
                (
                    "import sys\n"
                    "sys.stdout.buffer.write('{\"message\":\"姓 名：\"}'.encode('utf-8'))\n"
                ),
                encoding="utf-8",
            )

            result = run_optional("emit_utf8.py", project_root)

            self.assertEqual(result["status"], "ok")
            self.assertEqual(result["returncode"], 0)
            self.assertEqual(result["stdout"], '{"message":"姓 名："}')
            self.assertEqual(result["stderr"], "")

    def test_init_project_creates_default_templates_and_preview(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )

            self.assertEqual(result.returncode, 0, msg=result.stderr)
            self.assertTrue(
                (project_root / "templates" / "template.sample.docx").exists()
            )
            self.assertTrue(
                (project_root / "templates" / "template.user.docx").exists()
            )
            self.assertTrue((project_root / "out" / "preview.docx").exists())

            init_report = json.loads(
                (project_root / "logs" / "init_report.json").read_text(encoding="utf-8")
            )
            script_names = [item["script"] for item in init_report["script_results"]]
            self.assertIn("scan_template.py", script_names)
            self.assertIn("build_preview.py", script_names)

    def test_init_project_writes_language_preference_to_user_profile(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )

            self.assertEqual(result.returncode, 0, msg=result.stderr)

            user_profile = (project_root / "user" / "user.md").read_text(
                encoding="utf-8"
            )
            self.assertIn("- 语言：中文", user_profile)
            self.assertIn("- 语言偏好：zh-CN", user_profile)

    def test_init_project_uses_sanitized_default_template(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )

            self.assertEqual(result.returncode, 0, msg=result.stderr)
            template_doc = docx.Document(
                project_root / "templates" / "template.user.docx"
            )
            visible = [
                paragraph.text.strip()
                for paragraph in template_doc.paragraphs
                if paragraph.text.strip()
            ]

            self.assertIn("报告题目 / Report Title", visible)
            self.assertIn("姓 名：", visible)
            self.assertIn("学 号：", visible)
            self.assertIn("完成日期：", visible)

    def test_init_project_force_refreshes_default_template_assets(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            first_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(first_result.returncode, 0, msg=first_result.stderr)

            for relative in (
                "templates/template.sample.docx",
                "templates/template.user.docx",
            ):
                stale_doc = docx.Document()
                stale_doc.add_paragraph("stale template")
                stale_doc.save(project_root / relative)

            refresh_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                    "--force",
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(refresh_result.returncode, 0, msg=refresh_result.stderr)

            refreshed_doc = docx.Document(
                project_root / "templates" / "template.user.docx"
            )
            visible = [
                paragraph.text.strip()
                for paragraph in refreshed_doc.paragraphs
                if paragraph.text.strip()
            ]
            self.assertIn("报告题目 / Report Title", visible)

    def test_init_project_generates_template_recommendation(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            source_template = project_root / "voice-template.docx"
            write_style_poor_template(source_template)

            result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                    "--template",
                    str(source_template),
                ],
                capture_output=True,
                text=True,
            )

            self.assertEqual(result.returncode, 0, msg=result.stderr)
            self.assertTrue(
                (project_root / "templates" / "template.recommended.docx").exists()
            )
            recommendation = json.loads(
                (
                    project_root / "logs" / "template_style_recommendation.json"
                ).read_text(encoding="utf-8")
            )

            self.assertEqual(
                recommendation["recommended_template"],
                "./templates/template.recommended.docx",
            )
            self.assertIn("正文", recommendation["missing_styles"])
            self.assertIn("正文", recommendation["copied_styles"])
            self.assertTrue(recommendation["pending_acceptance"])

    def test_preview_summary_requires_confirmation_when_recommended_template_pending(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            source_template = project_root / "voice-template.docx"
            write_style_poor_template(source_template)

            result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                    "--template",
                    str(source_template),
                ],
                capture_output=True,
                text=True,
            )

            self.assertEqual(result.returncode, 0, msg=result.stderr)
            summary = json.loads(
                (project_root / "out" / "preview.summary.json").read_text(
                    encoding="utf-8"
                )
            )

            self.assertIn(
                "template style recommendation pending",
                summary["review"]["needs_confirmation"],
            )
            self.assertEqual(
                summary["template_recommendation"]["recommended_template"],
                "./templates/template.recommended.docx",
            )
            self.assertTrue(summary["template_recommendation"]["pending_acceptance"])

    def test_apply_template_recommendation_switches_primary_template(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            source_template = project_root / "voice-template.docx"
            write_style_poor_template(source_template)

            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                    "--template",
                    str(source_template),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            apply_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "recommend_template_styles.py"),
                    "--project-root",
                    str(project_root),
                    "--apply",
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(apply_result.returncode, 0, msg=apply_result.stderr)

            plan = json.loads(
                (project_root / "config" / "template.plan.json").read_text(
                    encoding="utf-8"
                )
            )
            self.assertEqual(
                plan["selection"]["primary_template"],
                "./templates/template.recommended.docx",
            )

            workflow = json.loads(
                (project_root / "workflow.json").read_text(encoding="utf-8")
            )
            self.assertEqual(
                workflow["templates"]["main_template"],
                "./templates/template.recommended.docx",
            )

            task_contract = yaml.safe_load(
                (project_root / "report.task.yaml").read_text(encoding="utf-8")
            )
            self.assertEqual(
                task_contract["inputs"]["template_path"],
                "./templates/template.recommended.docx",
            )

            user_template = docx.Document(
                project_root / "templates" / "template.user.docx"
            )
            user_style_names = {
                style.name
                for style in user_template.styles
                if getattr(style, "name", None)
            }
            self.assertNotIn("正文", user_style_names)

    def test_private_field_injection_builds_private_output(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            private_source = project_root.parent / "private-fields.json"
            private_source.write_text(
                json.dumps(
                    {"full_name": "Test User", "student_id": "S-001"},
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            inject_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "inject_private_fields.py"),
                    "--project-root",
                    str(project_root),
                    "--source",
                    str(private_source),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(inject_result.returncode, 0, msg=inject_result.stderr)

            private_doc = docx.Document(project_root / "out" / "private.docx")
            texts = [paragraph.text for paragraph in private_doc.paragraphs]
            self.assertTrue(any("Test User" in text for text in texts))
            self.assertTrue(any("S-001" in text for text in texts))

    def test_private_field_injection_preserves_cover_run_formatting(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            source_template = project_root / "styled-cover-template.docx"
            write_styled_cover_template(source_template)

            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                    "--template",
                    str(source_template),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            private_source = project_root.parent / "private-fields.json"
            private_source.write_text(
                json.dumps(
                    {"full_name": "田中 音声测试", "student_id": "VOICE-2026-0421"},
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            inject_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "inject_private_fields.py"),
                    "--project-root",
                    str(project_root),
                    "--source",
                    str(private_source),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(inject_result.returncode, 0, msg=inject_result.stderr)

            private_doc = docx.Document(project_root / "out" / "private.docx")
            paragraph = next(
                item
                for item in private_doc.paragraphs
                if "学 生 姓 名：" in item.text
            )

            self.assertEqual(paragraph.text, "学 生 姓 名：田中 音声测试")
            self.assertEqual(len(paragraph.runs), 2)
            self.assertEqual(run_font_settings(paragraph.runs[0])["ascii"], "Times New Roman")
            self.assertEqual(run_font_settings(paragraph.runs[0])["size"], "30")
            self.assertEqual(paragraph.runs[0].text, "学 生 姓 名：")
            self.assertEqual(paragraph.runs[1].text, "田中 音声测试")
            self.assertEqual(run_font_settings(paragraph.runs[1])["ascii"], "Times New Roman")
            self.assertEqual(run_font_settings(paragraph.runs[1])["eastAsia"], "仿宋")
            self.assertEqual(run_font_settings(paragraph.runs[1])["size"], "28")

    def test_build_preview_surfaces_unbound_cover_candidates_and_private_template(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            source_template = project_root / "unbound-cover-template.docx"
            write_cover_template_with_unbound_candidates(source_template)

            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                    "--template",
                    str(source_template),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            summary = json.loads(
                (project_root / "out" / "preview.summary.json").read_text(
                    encoding="utf-8"
                )
            )

            self.assertEqual(
                summary["field_binding"]["private_template"],
                {"full_name": "", "student_id": ""},
            )
            self.assertIn("学 院：", summary["field_binding"]["unbound_candidates"])
            self.assertIn("专 业：", summary["field_binding"]["unbound_candidates"])
            self.assertIn(
                "cover field candidates detected without bindings",
                summary["review"]["needs_confirmation"],
            )

    def test_scan_template_detects_cover_candidates_with_placeholder_suffixes(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            source_template = project_root / "placeholder-cover-template.docx"
            write_cover_template_with_placeholder_candidates(source_template)

            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                    "--template",
                    str(source_template),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            scan = json.loads(
                (project_root / "logs" / "template_scan.json").read_text(
                    encoding="utf-8"
                )
            )
            candidates = [item["text"] for item in scan["anchors"]["field_candidates"]]

            self.assertIn("学       院：", candidates)
            self.assertIn("专       业：", candidates)
            self.assertIn("评 阅 教 师：", candidates)
            self.assertIn("完 成 时 间：", candidates)

    def test_build_report_splits_markdown_into_multiple_paragraphs(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "# Summary\n\nFirst paragraph.\n\n## Details\n\nSecond paragraph.",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            texts = [
                paragraph.text.strip()
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip()
            ]

            self.assertIn("Summary", texts)
            self.assertIn("First paragraph.", texts)
            self.assertIn("Details", texts)
            self.assertIn("Second paragraph.", texts)

    def test_build_report_prefers_template_native_body_styles(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "# 报告题目\n\n## 一级标题\n\n### 二级标题\n\n正文段落示例。",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            rendered = {
                paragraph.text.strip(): paragraph.style.name
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip()
            }

            self.assertEqual(rendered["报告题目"], "题目")
            self.assertEqual(rendered["一级标题"], "标题2")
            self.assertEqual(rendered["二级标题"], "标题3")
            self.assertEqual(rendered["正文段落示例。"], "正文")

    def test_build_report_renders_fenced_code_block_as_single_cell_table(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Code Example\n\n```python\nprint('hello')\nprint('world')\n```",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)
            payload = json.loads(build_result.stdout)
            self.assertEqual(payload["code_blocks"]["styled"], 1)
            self.assertEqual(payload["code_blocks"]["highlighted"], 1)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            self.assertEqual(len(redacted_doc.tables), 1)
            table = redacted_doc.tables[0]
            self.assertEqual(len(table.rows), 1)
            self.assertEqual(len(table.columns), 1)
            self.assertEqual(table.cell(0, 0).paragraphs[0].text.strip(), "python")
            cell_text = "\n".join(
                paragraph.text for paragraph in table.cell(0, 0).paragraphs
            )
            self.assertIn("print('hello')", cell_text)
            self.assertIn("print('world')", cell_text)

    def test_build_report_highlights_c_code(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## C Example\n\n```c\nint main(void) {\n  return 0;\n}\n```",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)
            payload = json.loads(build_result.stdout)
            self.assertEqual(payload["code_blocks"]["highlighted"], 1)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            table = redacted_doc.tables[0]
            self.assertEqual(table.cell(0, 0).paragraphs[0].text.strip(), "c")

    def test_build_report_highlights_cpp_code(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## CPP Example\n\n```cpp\n#include <iostream>\nint main() {\n  std::cout << 1;\n}\n```",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)
            payload = json.loads(build_result.stdout)
            self.assertEqual(payload["code_blocks"]["highlighted"], 1)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            table = redacted_doc.tables[0]
            self.assertEqual(table.cell(0, 0).paragraphs[0].text.strip(), "cpp")

    def test_build_report_highlights_java_code(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Java Example\n\n```java\nclass Demo {\n  int x = 1;\n}\n```",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)
            payload = json.loads(build_result.stdout)
            self.assertEqual(payload["code_blocks"]["highlighted"], 1)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            table = redacted_doc.tables[0]
            self.assertEqual(table.cell(0, 0).paragraphs[0].text.strip(), "java")

    def test_build_report_requires_agent_handoff_for_unsupported_language(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                '## Rust Example\n\n```rust\nfn main() {\n    println!("hi");\n}\n```',
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(build_result.returncode, 0)
            payload = json.loads(build_result.stdout)
            self.assertEqual(payload["code_blocks"]["styled"], 1)
            self.assertEqual(payload["code_blocks"]["highlighted"], 0)
            self.assertEqual(len(payload["code_blocks"]["unsupported"]), 1)
            self.assertEqual(
                payload["code_blocks"]["unsupported"][0]["language"], "rust"
            )
            self.assertIsNone(payload["code_blocks"]["unsupported"][0]["normalized"])
            self.assertEqual(
                payload["code_blocks"]["unsupported"][0]["action"],
                "agent_handoff_required",
            )

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            table = redacted_doc.tables[0]
            self.assertEqual(table.cell(0, 0).paragraphs[0].text.strip(), "rust")

    def test_build_report_applies_code_theme_override(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "config" / "code-theme.user.json").write_text(
                json.dumps(
                    {
                        "base": "github-light",
                        "roles": {
                            "header_bg": "#EFEFEF",
                            "border": "#111111",
                            "header_fg": "#8A2BE2",
                            "keyword": "#AA0000",
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            (project_root / "docs" / "report_body.md").write_text(
                "## Code Example\n\n```python\nprint('hello')\n```",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)
            payload = json.loads(build_result.stdout)
            self.assertTrue(payload["code_blocks"]["theme"]["override_used"])
            self.assertEqual(payload["code_blocks"]["theme"]["name"], "github-light")
            self.assertEqual(payload["code_blocks"]["warnings"], [])

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            table = redacted_doc.tables[0]
            self.assertIn("EFEFEF", table.cell(0, 0).paragraphs[0]._p.xml.upper())
            self.assertIn("111111", table.cell(0, 0)._tc.xml.upper())

    def test_build_report_warns_on_invalid_code_theme_override(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "config" / "code-theme.user.json").write_text(
                json.dumps(
                    {
                        "base": "github-light",
                        "roles": {
                            "keyword": "not-a-hex",
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            (project_root / "docs" / "report_body.md").write_text(
                "## Code Example\n\n```python\nprint('hello')\n```",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)
            payload = json.loads(build_result.stdout)
            self.assertFalse(payload["code_blocks"]["theme"]["override_used"])
            self.assertTrue(payload["code_blocks"]["warnings"])

    def test_build_report_warns_on_non_object_code_theme_override(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "config" / "code-theme.user.json").write_text(
                "[]",
                encoding="utf-8",
            )
            (project_root / "docs" / "report_body.md").write_text(
                "## Code Example\n\n```python\nprint('hello')\n```",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)
            payload = json.loads(build_result.stdout)
            self.assertFalse(payload["code_blocks"]["theme"]["override_used"])
            self.assertTrue(payload["code_blocks"]["warnings"])

    def test_build_report_highlights_remaining_supported_languages_and_aliases(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                '## Languages\n\n```json\n{"a": 1}\n```\n\n```sh\necho hi\n```\n\n```yml\nname: demo\n```\n\n```sql\nselect 1;\n```\n\n```js\nconst n = 1;\n```\n\n```ts\nconst n: number = 1;\n```\n\n```c++\nint main() { return 0; }\n```',
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)
            payload = json.loads(build_result.stdout)
            self.assertEqual(payload["code_blocks"]["highlighted"], 7)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            headers = [
                table.cell(0, 0).paragraphs[0].text.strip()
                for table in redacted_doc.tables
            ]
            self.assertIn("json", headers)
            self.assertIn("bash", headers)
            self.assertIn("yaml", headers)
            self.assertIn("sql", headers)
            self.assertIn("javascript", headers)
            self.assertIn("typescript", headers)
            self.assertIn("cpp", headers)

    def test_build_report_renders_styled_plain_code_block(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Code Example\n\n```\nplain line 1\nplain line 2\n```",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            self.assertEqual(len(redacted_doc.tables), 1)
            table = redacted_doc.tables[0]
            cell_paragraphs = [
                paragraph.text.strip() for paragraph in table.cell(0, 0).paragraphs
            ]
            self.assertIn("Code", cell_paragraphs)
            self.assertIn("plain line 1", cell_paragraphs)
            self.assertIn("plain line 2", cell_paragraphs)

    def test_build_report_renders_markdown_lists(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Items\n\n- First bullet\n- Second bullet\n\n1. First number\n2. Second number\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            rendered = [
                (paragraph.text.strip(), paragraph.style.name)
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip()
            ]

            self.assertIn(("First bullet", "列表符号"), rendered)
            self.assertIn(("Second bullet", "列表符号"), rendered)
            self.assertIn(("First number", "列表编号"), rendered)
            self.assertIn(("Second number", "列表编号"), rendered)

    def test_build_report_restarts_ordered_lists_after_intervening_paragraph(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Items\n\n1. A\n2. B\n\nParagraph\n\n1. C\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_path = project_root / "out" / "redacted.docx"
            redacted_doc = docx.Document(redacted_path)
            ordered = {
                paragraph.text.strip(): paragraph_numbering_id(paragraph)
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip() in {"A", "B", "C"}
            }

            self.assertEqual(ordered["A"], ordered["B"])
            self.assertNotEqual(ordered["B"], ordered["C"])
            self.assertEqual(
                numbering_start_overrides(redacted_path)[ordered["C"]],
                1,
            )

    def test_build_report_restarts_ordered_lists_after_separate_list_block(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Items\n\n1. A\n2. B\n\n1. C\n2. D\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_path = project_root / "out" / "redacted.docx"
            redacted_doc = docx.Document(redacted_path)
            ordered = {
                paragraph.text.strip(): paragraph_numbering_id(paragraph)
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip() in {"A", "B", "C", "D"}
            }

            self.assertEqual(ordered["A"], ordered["B"])
            self.assertEqual(ordered["C"], ordered["D"])
            self.assertNotEqual(ordered["B"], ordered["C"])
            self.assertEqual(
                numbering_start_overrides(redacted_path)[ordered["C"]],
                1,
            )

    def test_build_report_preserves_custom_ordered_list_start(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Items\n\n3. C\n4. D\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_path = project_root / "out" / "redacted.docx"
            redacted_doc = docx.Document(redacted_path)
            ordered = {
                paragraph.text.strip(): paragraph_numbering_id(paragraph)
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip() in {"C", "D"}
            }

            self.assertEqual(ordered["C"], ordered["D"])
            self.assertEqual(
                numbering_start_overrides(redacted_path)[ordered["C"]],
                3,
            )

    def test_build_report_keeps_top_level_ordered_list_sequence_across_nested_items(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Items\n\n1. Parent A\n  1. Child one\n  2. Child two\n2. Parent B\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            ordered = {
                paragraph.text.strip(): paragraph_numbering_id(paragraph)
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip()
                in {"Parent A", "Parent B", "Child one", "Child two"}
            }

            self.assertEqual(ordered["Parent A"], ordered["Parent B"])
            self.assertEqual(ordered["Child one"], ordered["Child two"])
            self.assertNotEqual(ordered["Parent A"], ordered["Child one"])

    def test_build_report_lists_inherit_body_font_settings(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "正文段落示例。\n\n- First bullet\n\n1. First number\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            body_style = redacted_doc.styles["正文"]
            expected = style_font_settings(body_style)

            list_runs = [
                paragraph.runs[0]
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip() in {"First bullet", "First number"}
            ]

            self.assertEqual(len(list_runs), 2)
            for run in list_runs:
                self.assertEqual(run_font_settings(run), expected)

    def test_build_report_preserves_visible_list_markers_without_list_styles(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            source_template = project_root / "voice-template.docx"
            write_style_poor_template(source_template)
            strip_list_styles(source_template)

            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                    "--template",
                    str(source_template),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Items\n\n- First bullet\n- Second bullet\n\n1. First number\n2. Second number\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            rendered_texts = [
                paragraph.text.strip()
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip()
            ]

            self.assertIn("- First bullet", rendered_texts)
            self.assertIn("- Second bullet", rendered_texts)
            self.assertIn("1. First number", rendered_texts)
            self.assertIn("2. Second number", rendered_texts)

    def test_build_report_restarts_fallback_ordered_markers_without_list_styles(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            source_template = project_root / "voice-template.docx"
            write_style_poor_template(source_template)
            strip_list_styles(source_template)

            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                    "--template",
                    str(source_template),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Items\n\n1. A\n2. B\n\nParagraph\n\n1. C\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            rendered_texts = [
                paragraph.text.strip()
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip()
            ]

            self.assertIn("1. A", rendered_texts)
            self.assertIn("2. B", rendered_texts)
            self.assertIn("1. C", rendered_texts)

    def test_build_report_renders_simple_pipe_table(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Metrics\n\n| Name | Value |\n| --- | --- |\n| Alpha | 1 |\n| Beta | 2 |\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            self.assertEqual(len(redacted_doc.tables), 1)
            table = redacted_doc.tables[0]
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            self.assertEqual(rows, [["Name", "Value"], ["Alpha", "1"], ["Beta", "2"]])

    def test_build_report_inserts_existing_image(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            image_path = project_root / "docs" / "images" / "arch.png"
            image_path.parent.mkdir(parents=True, exist_ok=True)
            image_path.write_bytes(TEST_PNG_BYTES)

            (project_root / "docs" / "report_body.md").write_text(
                "## Figures\n\n![Architecture](images/arch.png)\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)
            payload = json.loads(build_result.stdout)
            self.assertEqual(payload["images"]["failed"], [])
            self.assertEqual(len(payload["images"]["inserted"]), 1)
            self.assertEqual(
                payload["images"]["inserted"][0]["path"], "images/arch.png"
            )

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            self.assertTrue(
                any(
                    "<w:drawing" in paragraph._p.xml
                    for paragraph in redacted_doc.paragraphs
                )
            )

    def test_build_report_renders_centered_figure_with_caption_below(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            image_path = project_root / "docs" / "images" / "arch.png"
            image_path.parent.mkdir(parents=True, exist_ok=True)
            image_path.write_bytes(TEST_PNG_BYTES)

            (project_root / "docs" / "report_body.md").write_text(
                "## Figures\n\n![Architecture](images/arch.png)\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            image_indexes = [
                index
                for index, paragraph in enumerate(redacted_doc.paragraphs)
                if "<w:drawing" in paragraph._p.xml
            ]
            self.assertEqual(len(image_indexes), 1)

            image_paragraph = redacted_doc.paragraphs[image_indexes[0]]
            self.assertEqual(image_paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertIn("wp:inline", image_paragraph._p.xml)
            self.assertNotIn("wp:anchor", image_paragraph._p.xml)
            self.assertNotIn("wrapTopAndBottom", image_paragraph._p.xml)

            caption_paragraph = redacted_doc.paragraphs[image_indexes[0] + 1]
            self.assertEqual(caption_paragraph.style.name, "图题")
            self.assertTrue(caption_paragraph.text.strip().startswith("图1"))

            with zipfile.ZipFile(
                project_root / "out" / "redacted.docx", "r"
            ) as docx_zip:
                document_xml = docx_zip.read("word/document.xml").decode("utf-8")
            self.assertIn("<wp:inline", document_xml)
            self.assertNotIn("<wp:anchor", document_xml)

    def test_build_report_renders_centered_table_with_caption_and_cell_formatting(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Metrics\n\n| Name | Value |\n| --- | --- |\n| Alpha | 1 |\n| Beta | 2 |\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            self.assertEqual(len(redacted_doc.tables), 1)
            table = redacted_doc.tables[0]
            self.assertEqual(table.alignment, WD_TABLE_ALIGNMENT.CENTER)

            caption_element = table._tbl.getprevious()
            self.assertIsNotNone(caption_element)
            caption_xml = caption_element.xml
            self.assertIn("SEQ 表", caption_xml)
            self.assertIn("bookmarkStart", caption_xml)

            caption_texts = [
                (paragraph.text.strip(), paragraph.style.name)
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip()
            ]
            self.assertIn(("表1 Metrics", "表题"), caption_texts)

            cell_paragraph = table.cell(0, 0).paragraphs[0]
            self.assertEqual(cell_paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(
                table.cell(0, 0).vertical_alignment,
                WD_CELL_VERTICAL_ALIGNMENT.CENTER,
            )
            self.assertIsNotNone(cell_paragraph.paragraph_format.first_line_indent)
            self.assertEqual(cell_paragraph.paragraph_format.first_line_indent.pt, 0.0)
            self.assertEqual(cell_paragraph.paragraph_format.line_spacing, 1.5)

            cell_run = cell_paragraph.runs[0]
            self.assertEqual(cell_run.font.name, "宋体")
            self.assertIsNotNone(cell_run.font.size)
            self.assertEqual(cell_run.font.size.pt, 10.5)
            self.assertEqual(run_font_settings(cell_run)["eastAsia"], "宋体")

    def test_build_report_strips_section_number_from_table_caption(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## 3 实验环境与参数\n\n| 参数 | 数值 |\n| --- | --- |\n| 采样率 | 16000 Hz |\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            caption_texts = [
                paragraph.text.strip()
                for paragraph in redacted_doc.paragraphs
                if paragraph.style.name == "表题" and paragraph.text.strip()
            ]

            self.assertIn("表1 实验环境与参数", caption_texts)
            self.assertNotIn("表1 3 实验环境与参数", caption_texts)

    def test_build_report_replaces_figure_and_table_reference_placeholders(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            image_path = project_root / "docs" / "images" / "arch.png"
            image_path.parent.mkdir(parents=True, exist_ok=True)
            image_path.write_bytes(TEST_PNG_BYTES)

            plan_path = project_root / "config" / "template.plan.json"
            plan = json.loads(plan_path.read_text(encoding="utf-8"))
            plan["semantics"]["cross_references"]["figure_table_enabled"] = True
            plan_path.write_text(
                json.dumps(plan, ensure_ascii=False, indent=2) + "\n",
                encoding="utf-8",
            )

            (project_root / "docs" / "report_body.md").write_text(
                "## Figures\n\n![Architecture](images/arch.png)\n\n[[REF:figure:fig_0001|见下图]] 展示了系统结构。\n\n## Metrics\n\n| Name | Value |\n| --- | --- |\n| Alpha | 1 |\n\n[[REF:table:tbl_0001|见上表]] 汇总了实验结果。\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            rendered = {
                paragraph.text.strip(): paragraph._p.xml
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip()
            }

            self.assertIn("见下图1 展示了系统结构。", rendered)
            self.assertIn('w:anchor="fig_0001"', rendered["见下图1 展示了系统结构。"])
            self.assertIn("见上表1 汇总了实验结果。", rendered)
            self.assertIn('w:anchor="tbl_0001"', rendered["见上表1 汇总了实验结果。"])

    def test_build_report_applies_reference_style_in_reference_section(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## 参考文献\n\n[1] 作者. 题名[J]. 期刊名, 2024, 1(1): 1-10.\n\n[2] 作者. 书名[M]. 北京: 出版社, 2023.\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            rendered = {
                paragraph.text.strip(): paragraph.style.name
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip().startswith("[")
            }

            self.assertEqual(
                rendered["[1] 作者. 题名[J]. 期刊名, 2024, 1(1): 1-10."],
                "参考文献",
            )
            self.assertEqual(
                rendered["[2] 作者. 书名[M]. 北京: 出版社, 2023."],
                "参考文献",
            )

    def test_build_report_applies_reference_style_in_numbered_reference_section(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## 二、参考文献\n\n[1] 作者. 题名[J]. 期刊名, 2024, 1(1): 1-10.\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            rendered = {
                paragraph.text.strip(): paragraph.style.name
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip().startswith("[")
            }
            self.assertEqual(
                rendered["[1] 作者. 题名[J]. 期刊名, 2024, 1(1): 1-10."],
                "参考文献",
            )

    def test_build_report_applies_reference_style_to_numbered_reference_entries(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## 参考文献\n\n1. Author. Title[J]. Journal, 2024, 1(1): 1-10.\n2. Author. Book[M]. Beijing: Press, 2023.\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            rendered = {
                paragraph.text.strip(): paragraph.style.name
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip().startswith(("Author. Title", "Author. Book"))
            }
            self.assertEqual(
                rendered["Author. Title[J]. Journal, 2024, 1(1): 1-10."],
                "参考文献",
            )
            self.assertEqual(
                rendered["Author. Book[M]. Beijing: Press, 2023."],
                "参考文献",
            )

    def test_reference_citation_section_does_not_reuse_reference_entry_style(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## 参考文献引用测试\n\n本节专门覆盖文献交叉引用，不应继承参考文献条目样式。\n\n## 参考文献\n\n[1] 作者. 题名[J]. 期刊名, 2024, 1(1): 1-10.\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(build_result.returncode, 0, msg=build_result.stderr)

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            citation_paragraph = next(
                paragraph
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip()
                == "本节专门覆盖文献交叉引用，不应继承参考文献条目样式。"
            )
            reference_paragraph = next(
                paragraph
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip().startswith("[1] 作者.")
            )

            self.assertNotEqual(citation_paragraph.style.name, "参考文献")
            self.assertIn(citation_paragraph.style.name, {"正文", "正文1", "Normal"})
            self.assertIsNone(citation_paragraph.paragraph_format.left_indent)
            self.assertIsNone(citation_paragraph.paragraph_format.first_line_indent)
            self.assertNotEqual(
                citation_paragraph.style.name,
                reference_paragraph.style.name,
            )

    def test_build_report_reports_failed_image_insertions(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            (project_root / "docs" / "report_body.md").write_text(
                "## Figures\n\n![Missing](images/missing.png)\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(build_result.returncode, 0)
            payload = json.loads(build_result.stdout)
            self.assertEqual(
                payload["redacted"], str(project_root / "out" / "redacted.docx")
            )
            self.assertEqual(payload["images"]["inserted"], [])
            self.assertEqual(len(payload["images"]["failed"]), 1)
            self.assertEqual(
                payload["images"]["failed"][0]["path"], "images/missing.png"
            )

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            rendered = [
                paragraph.text.strip()
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip()
            ]
            self.assertTrue(
                any(
                    text.startswith(
                        "[Image Insert Failed] Missing (images/missing.png):"
                    )
                    for text in rendered
                )
            )

    def test_build_report_reports_corrupt_image_insertions(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            project_root = Path(temp_dir)
            init_result = subprocess.run(
                [
                    str(PYTHON),
                    str(PROJECT_ROOT / "scripts" / "init_project.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

            image_path = project_root / "docs" / "images" / "bad.png"
            image_path.parent.mkdir(parents=True, exist_ok=True)
            image_path.write_bytes(b"not-a-real-png")

            (project_root / "docs" / "report_body.md").write_text(
                "## Figures\n\n![Broken](images/bad.png)\n",
                encoding="utf-8",
            )

            build_result = subprocess.run(
                [
                    str(PYTHON),
                    str(project_root / "scripts" / "build_report.py"),
                    "--project-root",
                    str(project_root),
                ],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(build_result.returncode, 0)
            payload = json.loads(build_result.stdout)
            self.assertEqual(payload["images"]["inserted"], [])
            self.assertEqual(len(payload["images"]["failed"]), 1)
            self.assertEqual(payload["images"]["failed"][0]["path"], "images/bad.png")
            self.assertNotEqual(
                payload["images"]["failed"][0]["reason"], "file not found"
            )

            redacted_doc = docx.Document(project_root / "out" / "redacted.docx")
            rendered = [
                paragraph.text.strip()
                for paragraph in redacted_doc.paragraphs
                if paragraph.text.strip()
            ]
            self.assertTrue(
                any(
                    text.startswith("[Image Insert Failed] Broken (images/bad.png):")
                    for text in rendered
                )
            )


if __name__ == "__main__":
    unittest.main()
