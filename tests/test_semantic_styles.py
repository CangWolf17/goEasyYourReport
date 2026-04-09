from __future__ import annotations

import json
import re
import shutil
import subprocess
import unittest
import uuid
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

import docx

from scripts._docx_integrity import validate_docx_package


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PYTHON = Path(r"D:\Miniconda\python.exe")
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
STYLE_XML_DECLARATION_PATTERN = re.compile(rb"^<\?xml[^?]*\?>")
STYLE_XML_ROOT_PATTERN = re.compile(rb"<w:styles\b[^>]*>")
STYLE_XML_NAMESPACES = {
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": W_NS,
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
    "w16": "http://schemas.microsoft.com/office/word/2018/wordml",
    "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
    "w16cid": "http://schemas.microsoft.com/office/word/2016/wordml/cid",
    "w16sdtdh": "http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash",
    "w16se": "http://schemas.microsoft.com/office/word/2015/wordml/symex",
}


def write_template(
    path: Path,
    *,
    include_toc: bool = False,
    include_reference_block: bool = True,
) -> None:
    document = docx.Document()
    document.add_paragraph("课程考核报告")
    document.add_paragraph("姓 名：")
    document.add_paragraph("学 号：")
    document.add_paragraph("完成日期：")
    if include_toc:
        toc = document.add_paragraph("目录")
        toc.style = "TOC Heading"
    document.add_heading("一、正文内容", level=1)
    document.add_paragraph("这里是正文。")
    if include_reference_block:
        document.add_heading("二、参考文献", level=1)
        document.add_paragraph("[1] Reference")
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def style_outline_level(docx_path: Path, style_name: str) -> int | None:
    with zipfile.ZipFile(docx_path, "r") as docx_zip:
        root = ET.fromstring(docx_zip.read("word/styles.xml"))
    for style in root.findall(f"{{{W_NS}}}style"):
        name = style.find(f"{{{W_NS}}}name")
        if name is None or name.get(f"{{{W_NS}}}val") != style_name:
            continue
        p_pr = style.find(f"{{{W_NS}}}pPr")
        if p_pr is None:
            return None
        outline = p_pr.find(f"{{{W_NS}}}outlineLvl")
        if outline is None:
            return None
        raw = outline.get(f"{{{W_NS}}}val")
        return None if raw is None else int(raw)
    return None


def style_id(docx_path: Path, style_name: str) -> str | None:
    with zipfile.ZipFile(docx_path, "r") as docx_zip:
        root = ET.fromstring(docx_zip.read("word/styles.xml"))
    for style in root.findall(f"{{{W_NS}}}style"):
        name = style.find(f"{{{W_NS}}}name")
        if name is None or name.get(f"{{{W_NS}}}val") != style_name:
            continue
        return style.get(f"{{{W_NS}}}styleId") or style.get("styleId")
    return None


def style_dependency_value(docx_path: Path, style_name: str, dependency: str) -> str | None:
    with zipfile.ZipFile(docx_path, "r") as docx_zip:
        root = ET.fromstring(docx_zip.read("word/styles.xml"))
    for style in root.findall(f"{{{W_NS}}}style"):
        name = style.find(f"{{{W_NS}}}name")
        if name is None or name.get(f"{{{W_NS}}}val") != style_name:
            continue
        dependency_element = style.find(f"{{{W_NS}}}{dependency}")
        if dependency_element is None:
            return None
        return dependency_element.get(f"{{{W_NS}}}val")
    return None


def rewrite_normal_style_id(docx_path: Path, new_style_id: str) -> None:
    with zipfile.ZipFile(docx_path, "r") as source_zip:
        entries = {
            info.filename: source_zip.read(info.filename)
            for info in source_zip.infolist()
        }

    original_styles_xml = entries["word/styles.xml"]
    styles_root = ET.fromstring(original_styles_xml)
    old_style_id: str | None = None
    for style in styles_root.findall(f"{{{W_NS}}}style"):
        name = style.find(f"{{{W_NS}}}name")
        if name is None or name.get(f"{{{W_NS}}}val") != "Normal":
            continue
        old_style_id = style.get(f"{{{W_NS}}}styleId") or style.get("styleId")
        style.set(f"{{{W_NS}}}styleId", new_style_id)
        break

    if old_style_id is None:
        raise AssertionError("Normal style not found in styles.xml")

    for style in styles_root.findall(f"{{{W_NS}}}style"):
        for dependency in ("basedOn", "next", "link"):
            dependency_element = style.find(f"{{{W_NS}}}{dependency}")
            if dependency_element is None:
                continue
            if dependency_element.get(f"{{{W_NS}}}val") == old_style_id:
                dependency_element.set(f"{{{W_NS}}}val", new_style_id)

    for prefix, uri in STYLE_XML_NAMESPACES.items():
        ET.register_namespace(prefix, uri)

    serialized = ET.tostring(
        styles_root,
        encoding="utf-8",
        xml_declaration=True,
    )
    original_declaration = STYLE_XML_DECLARATION_PATTERN.search(original_styles_xml)
    if original_declaration is not None:
        serialized = STYLE_XML_DECLARATION_PATTERN.sub(
            original_declaration.group(0),
            serialized,
            count=1,
        )
    original_root = STYLE_XML_ROOT_PATTERN.search(original_styles_xml)
    serialized_root = STYLE_XML_ROOT_PATTERN.search(serialized)
    if original_root is not None and serialized_root is not None:
        serialized = (
            serialized[: serialized_root.start()]
            + original_root.group(0)
            + serialized[serialized_root.end() :]
        )

    entries["word/styles.xml"] = serialized
    with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as output_zip:
        for filename, content in entries.items():
            output_zip.writestr(filename, content)


def cell_border_values(cell) -> dict[str, str]:
    tc_pr = cell._tc.find(f"{{{W_NS}}}tcPr")
    if tc_pr is None:
        return {}
    borders = tc_pr.find(f"{{{W_NS}}}tcBorders")
    if borders is None:
        return {}
    values: dict[str, str] = {}
    for edge in ("top", "bottom"):
        element = borders.find(f"{{{W_NS}}}{edge}")
        if element is not None:
            values[edge] = element.get(f"{{{W_NS}}}val", "")
    return values


class SemanticStyleTests(unittest.TestCase):
    def create_project(self) -> Path:
        sandbox_root = PROJECT_ROOT / "temp" / "semantic-style-tests"
        sandbox_root.mkdir(parents=True, exist_ok=True)
        project_root = sandbox_root / uuid.uuid4().hex
        project_root.mkdir(parents=True, exist_ok=True)
        self.addCleanup(lambda: shutil.rmtree(project_root, ignore_errors=True))
        result = subprocess.run(
            [
                str(PYTHON),
                str(PROJECT_ROOT / "scripts" / "init_project.py"),
                "--project-root",
                str(project_root),
            ],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        return project_root

    def run_json(
        self, project_root: Path, script_name: str, *extra_args: str
    ) -> dict[str, object]:
        result = subprocess.run(
            [
                str(PYTHON),
                str(project_root / "scripts" / script_name),
                "--project-root",
                str(project_root),
                *extra_args,
            ],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        return json.loads(result.stdout)

    def run_completed(
        self, project_root: Path, script_name: str, *extra_args: str
    ) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [
                str(PYTHON),
                str(project_root / "scripts" / script_name),
                "--project-root",
                str(project_root),
                *extra_args,
            ],
            capture_output=True,
            text=True,
        )

    def write_user_template(
        self,
        project_root: Path,
        *,
        include_toc: bool = False,
        include_reference_block: bool = True,
    ) -> None:
        write_template(
            project_root / "templates" / "template.user.docx",
            include_toc=include_toc,
            include_reference_block=include_reference_block,
        )

    def remove_reference_block(self, project_root: Path) -> None:
        template_path = project_root / "templates" / "template.user.docx"
        document = docx.Document(template_path)
        for paragraph in list(document.paragraphs):
            if paragraph.text.strip() in {"二、参考文献", "[1] 参考文献条目示例"}:
                paragraph._element.getparent().remove(paragraph._element)
        document.save(template_path)

    def test_scan_template_reports_semantic_style_candidates(self) -> None:
        project_root = self.create_project()

        scan = self.run_json(project_root, "scan_template.py")

        self.assertIn("styles", scan)
        self.assertIn("题目", scan["styles"]["available"])
        self.assertIn("标题2", scan["styles"]["heading_candidates"])
        self.assertIn("标题3", scan["styles"]["heading_candidates"])
        self.assertIn("正文", scan["styles"]["body_candidates"])

    def test_scan_template_reports_toc_signal(self) -> None:
        project_root = self.create_project()
        self.write_user_template(project_root, include_toc=True)

        scan = self.run_json(project_root, "scan_template.py")

        self.assertEqual(scan["toc"]["detected"], True)
        self.assertEqual(scan["toc"]["kind"], "placeholder")

    def test_scan_template_reports_reference_block_signal(self) -> None:
        project_root = self.create_project()
        self.write_user_template(project_root, include_reference_block=False)

        scan = self.run_json(project_root, "scan_template.py")

        self.assertEqual(scan["reference_block"]["present"], False)

    def test_scan_template_detects_reference_block_with_trailing_colon(self) -> None:
        project_root = self.create_project()
        template_path = project_root / "templates" / "template.user.docx"
        document = docx.Document(template_path)
        heading = next(
            paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "二、参考文献"
        )
        heading.text = "二、参考文献："
        document.save(template_path)

        scan = self.run_json(project_root, "scan_template.py")

        self.assertEqual(scan["reference_block"]["present"], True)

    def test_build_preview_includes_semantic_confirmation_payload(self) -> None:
        project_root = self.create_project()
        self.write_user_template(project_root, include_toc=True)
        self.run_json(project_root, "scan_template.py")

        preview = self.run_json(project_root, "build_preview.py")
        summary = json.loads(Path(preview["summary"]).read_text(encoding="utf-8"))

        self.assertIn("semantics", summary)
        self.assertIn("style_candidates", summary["semantics"])
        self.assertIn("style_gaps", summary["semantics"])
        self.assertEqual(summary["semantics"]["toc"]["detected"], True)
        self.assertIn(
            "template outline semantics incomplete",
            summary["review"]["needs_confirmation"],
        )
        self.assertIn(
            "list style semantics unresolved",
            summary["review"]["needs_confirmation"],
        )
        self.assertIn(
            "toc detected; confirm whether to enable",
            summary["review"]["needs_confirmation"],
        )
        self.assertIn(
            "confirm whether to insert figure/table cross references",
            summary["review"]["needs_confirmation"],
        )

    def test_default_template_contains_semantic_styles(self) -> None:
        project_root = self.create_project()
        document = docx.Document(project_root / "templates" / "template.user.docx")
        style_names = {
            style.name for style in document.styles if getattr(style, "name", None)
        }

        for expected in (
            "题目",
            "标题2",
            "标题3",
            "标题4",
            "正文",
            "图题",
            "表题",
            "参考文献",
            "列表编号",
            "列表符号",
        ):
            self.assertIn(expected, style_names)

    def test_default_template_sets_expected_outline_levels(self) -> None:
        project_root = self.create_project()
        template_path = project_root / "templates" / "template.user.docx"

        self.assertIsNone(style_outline_level(template_path, "题目"))
        self.assertEqual(style_outline_level(template_path, "标题2"), 0)
        self.assertEqual(style_outline_level(template_path, "标题3"), 1)
        self.assertEqual(style_outline_level(template_path, "标题4"), 2)

    def test_recommended_template_backfills_semantic_styles_and_outline_levels(
        self,
    ) -> None:
        project_root = self.create_project()
        source_template = project_root / "voice-template.docx"
        write_template(source_template, include_reference_block=False)

        init_result = self.run_completed(project_root, "init_project.py", "--template", str(source_template), "--force")
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        recommendation = self.run_json(project_root, "recommend_template_styles.py")
        recommended = project_root / recommendation["recommended_template"].replace("./", "")
        document = docx.Document(recommended)
        style_names = {
            style.name for style in document.styles if getattr(style, "name", None)
        }

        self.assertIn("题目", style_names)
        self.assertIn("标题2", style_names)
        self.assertIn("列表编号", style_names)
        self.assertIn("列表符号", style_names)
        self.assertEqual(style_outline_level(recommended, "标题2"), 0)
        self.assertEqual(style_outline_level(recommended, "标题3"), 1)
        self.assertEqual(style_outline_level(recommended, "标题4"), 2)

    def test_recommended_template_falls_back_to_repo_donor_when_project_sample_is_stale(
        self,
    ) -> None:
        project_root = self.create_project()
        source_template = project_root / "voice-template.docx"
        write_template(source_template, include_reference_block=False)

        init_result = self.run_completed(
            project_root,
            "init_project.py",
            "--template",
            str(source_template),
            "--force",
        )
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        stale_sample = project_root / "templates" / "template.sample.docx"
        write_template(stale_sample, include_reference_block=False)

        recommendation = self.run_json(project_root, "recommend_template_styles.py")
        recommended = project_root / recommendation["recommended_template"].replace("./", "")
        document = docx.Document(recommended)
        style_names = {
            style.name for style in document.styles if getattr(style, "name", None)
        }

        self.assertIn("列表编号", style_names)
        self.assertIn("列表符号", style_names)
        self.assertEqual(style_outline_level(recommended, "标题2"), 0)
        self.assertEqual(style_outline_level(recommended, "标题3"), 1)
        self.assertEqual(style_outline_level(recommended, "标题4"), 2)

    def test_recommended_template_remaps_style_dependencies_to_target_style_ids(
        self,
    ) -> None:
        project_root = self.create_project()
        source_template = project_root / "voice-template.docx"
        write_template(source_template, include_reference_block=False)

        init_result = self.run_completed(
            project_root,
            "init_project.py",
            "--template",
            str(source_template),
            "--force",
        )
        self.assertEqual(init_result.returncode, 0, msg=init_result.stderr)

        user_template = project_root / "templates" / "template.user.docx"
        rewrite_normal_style_id(user_template, "a")
        self.assertEqual(style_id(user_template, "Normal"), "a")

        recommendation = self.run_json(project_root, "recommend_template_styles.py")
        recommended = project_root / recommendation["recommended_template"].replace("./", "")
        normal_style_id = style_id(recommended, "Normal")

        self.assertEqual(normal_style_id, "a")
        for style_name in ("题目", "标题2", "标题3", "正文", "图题", "表题", "参考文献"):
            self.assertEqual(
                style_dependency_value(recommended, style_name, "basedOn"),
                normal_style_id,
                msg=style_name,
            )

        report = validate_docx_package(recommended)
        self.assertTrue(report["ok"], msg=report["errors"])

    def test_rendered_headings_use_semantic_word_styles(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "# 报告题目\n\n## 一级标题\n\n### 二级标题\n\n正文段落。",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        rendered = {
            paragraph.text.strip(): paragraph.style.name
            for paragraph in redacted.paragraphs
            if paragraph.text.strip()
        }

        self.assertEqual(rendered["报告题目"], "题目")
        self.assertEqual(rendered["一级标题"], "标题2")
        self.assertEqual(rendered["二级标题"], "标题3")
        self.assertEqual(style_outline_level(project_root / "out" / "redacted.docx", "标题2"), 0)
        self.assertEqual(style_outline_level(project_root / "out" / "redacted.docx", "标题3"), 1)

    def test_document_title_does_not_enter_outline_levels(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "# 报告题目\n\n## 一级标题",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        self.assertIsNone(style_outline_level(project_root / "out" / "redacted.docx", "题目"))
        self.assertEqual(style_outline_level(project_root / "out" / "redacted.docx", "标题2"), 0)

    def test_lists_use_word_list_styles_when_semantics_available(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "- First bullet\n- Second bullet\n\n1. First number\n2. Second number\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        rendered = {
            paragraph.text.strip(): paragraph.style.name
            for paragraph in redacted.paragraphs
            if paragraph.text.strip()
        }

        self.assertEqual(rendered["First bullet"], "列表符号")
        self.assertEqual(rendered["Second bullet"], "列表符号")
        self.assertEqual(rendered["First number"], "列表编号")
        self.assertEqual(rendered["Second number"], "列表编号")

    def test_reference_entries_use_reference_style_only_when_reference_block_exists(
        self,
    ) -> None:
        with_reference = self.create_project()
        (with_reference / "docs" / "report_body.md").write_text(
            "## 参考文献\n\n[1] 作者. 题名[J]. 期刊名, 2024.\n",
            encoding="utf-8",
        )
        result = self.run_completed(with_reference, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(with_reference / "out" / "redacted.docx")
        rendered = {
            paragraph.text.strip(): paragraph.style.name
            for paragraph in redacted.paragraphs
            if paragraph.text.strip().startswith("[1]")
        }
        self.assertEqual(rendered["[1] 作者. 题名[J]. 期刊名, 2024."], "参考文献")

        without_reference = self.create_project()
        self.remove_reference_block(without_reference)
        self.run_json(without_reference, "scan_template.py")
        (without_reference / "docs" / "report_body.md").write_text(
            "## 参考文献\n\n[1] 作者. 题名[J]. 期刊名, 2024.\n",
            encoding="utf-8",
        )
        result = self.run_completed(without_reference, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(without_reference / "out" / "redacted.docx")
        rendered = {
            paragraph.text.strip(): paragraph.style.name
            for paragraph in redacted.paragraphs
            if paragraph.text.strip().startswith("[1]")
        }
        self.assertNotEqual(rendered["[1] 作者. 题名[J]. 期刊名, 2024."], "参考文献")

    def test_default_table_border_policy_matches_repo_rule(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "| 参数 | 数值 |\n| --- | --- |\n| 采样率 | 16000 |\n| 时长 | 1.0 s |\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        table = redacted.tables[0]

        first_row = table.rows[0].cells[0]
        middle_row = table.rows[1].cells[0]
        last_row = table.rows[2].cells[0]

        self.assertEqual(cell_border_values(first_row).get("top"), "single")
        self.assertEqual(cell_border_values(first_row).get("bottom"), "single")
        self.assertNotIn("top", cell_border_values(middle_row))
        self.assertNotIn("bottom", cell_border_values(middle_row))
        self.assertEqual(cell_border_values(last_row).get("bottom"), "single")
        self.assertIn('insideH w:val="nil"', table._tbl.xml)
        self.assertIn('insideV w:val="nil"', table._tbl.xml)
        self.assertNotIn("tblStyle", table._tbl.xml)

    def test_first_column_bold_heuristic_only_applies_to_row_labels(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "## Labels\n\n| 参数 | 数值 |\n| --- | --- |\n| 采样率 | 16000 |\n| 时长 | 1.0 s |\n\n## Numbers\n\n| 序号 | 数值 |\n| --- | --- |\n| 1 | A |\n| 2 | B |\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        first_table = redacted.tables[0]
        second_table = redacted.tables[1]

        self.assertTrue(all(run.bold for run in first_table.cell(1, 0).paragraphs[0].runs if run.text))
        self.assertTrue(all(run.bold for run in first_table.cell(2, 0).paragraphs[0].runs if run.text))
        self.assertFalse(any(run.bold for run in second_table.cell(1, 0).paragraphs[0].runs if run.text))
        self.assertFalse(any(run.bold for run in second_table.cell(2, 0).paragraphs[0].runs if run.text))


if __name__ == "__main__":
    unittest.main()
