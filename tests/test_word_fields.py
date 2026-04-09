from __future__ import annotations

import base64
import json
import shutil
import subprocess
import sys
import unittest
import uuid
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

from scripts._docx_xml import insert_paragraph_after


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PYTHON = Path(r"D:\Miniconda\python.exe")
WINWORD = Path(r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE")
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TEST_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
)


class WordFieldTests(unittest.TestCase):
    def create_project(self) -> Path:
        sandbox_root = PROJECT_ROOT / "temp" / "word-field-tests"
        sandbox_root.mkdir(parents=True, exist_ok=True)
        project_root = sandbox_root / uuid.uuid4().hex
        project_root.mkdir(parents=True, exist_ok=True)
        self.addCleanup(lambda: shutil.rmtree(project_root, ignore_errors=True))
        result = subprocess.run(
            [
                str(PYTHON),
                str(PROJECT_ROOT / "scripts" / "init_project.py"),
                "--project-root",
                str(project_root),
            ],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        return project_root

    def run_completed(
        self, project_root: Path, script_name: str, *extra_args: str
    ) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [
                str(PYTHON),
                str(project_root / "scripts" / script_name),
                "--project-root",
                str(project_root),
                *extra_args,
            ],
            capture_output=True,
            text=True,
        )

    def insert_toc_placeholder(self, project_root: Path) -> None:
        template_path = project_root / "templates" / "template.user.docx"
        template = docx.Document(template_path)
        anchor = next(
            paragraph
            for paragraph in template.paragraphs
            if getattr(paragraph.style, "name", "") == "标题2"
        )
        anchor.insert_paragraph_before("目录")
        template.save(template_path)

        scan_result = self.run_completed(project_root, "scan_template.py")
        self.assertEqual(scan_result.returncode, 0, msg=scan_result.stderr)

    def set_toc_confirmation(
        self, project_root: Path, *, enabled: bool, needs_confirmation: bool
    ) -> None:
        plan_path = project_root / "config" / "template.plan.json"
        plan = json.loads(plan_path.read_text(encoding="utf-8"))
        toc = plan.setdefault("semantics", {}).setdefault("toc", {})
        toc["detected"] = True
        toc["kind"] = "placeholder"
        toc["enabled"] = enabled
        toc["needs_confirmation"] = needs_confirmation
        toc["source"] = "template_scan"
        plan_path.write_text(
            json.dumps(plan, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

    def set_cross_reference_confirmation(
        self, project_root: Path, enabled: bool | str
    ) -> None:
        plan_path = project_root / "config" / "template.plan.json"
        plan = json.loads(plan_path.read_text(encoding="utf-8"))
        cross_references = (
            plan.setdefault("semantics", {}).setdefault("cross_references", {})
        )
        cross_references["mode"] = "postprocess"
        cross_references["figure_table_enabled"] = enabled
        plan_path.write_text(
            json.dumps(plan, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

    def add_extra_cover_page_break(self, project_root: Path) -> None:
        from docx.enum.text import WD_BREAK

        template_path = project_root / "templates" / "template.user.docx"
        template = docx.Document(template_path)
        anchor = next(
            paragraph
            for paragraph in template.paragraphs
            if getattr(paragraph.style, "name", "") in {"Heading 1", "标题2"}
        )
        extra_break = anchor.insert_paragraph_before()
        extra_break.add_run().add_break(WD_BREAK.PAGE)
        template.save(template_path)

        scan_result = self.run_completed(project_root, "scan_template.py")
        self.assertEqual(scan_result.returncode, 0, msg=scan_result.stderr)

    def write_report_body(self, project_root: Path, content: str) -> None:
        (project_root / "docs" / "report_body.md").write_text(
            content,
            encoding="utf-8",
        )

    def assert_toc_style_matches_policy(self, style) -> None:
        self.assertEqual(style.font.name, "宋体")
        self.assertIsNotNone(style.font.size)
        self.assertEqual(style.font.size.pt, 14.0)
        self.assertEqual(style.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(style.paragraph_format.line_spacing, 1.5)
        self.assertIsNotNone(style.paragraph_format.left_indent)
        self.assertEqual(style.paragraph_format.left_indent.pt, 0.0)
        self.assertIsNotNone(style.paragraph_format.first_line_indent)
        self.assertEqual(style.paragraph_format.first_line_indent.pt, 0.0)

        r_pr = style.element.find(f"{{{W_NS}}}rPr")
        self.assertIsNotNone(r_pr)
        r_fonts = r_pr.find(f"{{{W_NS}}}rFonts")
        self.assertIsNotNone(r_fonts)
        self.assertEqual(r_fonts.get(f"{{{W_NS}}}ascii"), "宋体")
        self.assertEqual(r_fonts.get(f"{{{W_NS}}}hAnsi"), "宋体")
        self.assertEqual(r_fonts.get(f"{{{W_NS}}}eastAsia"), "宋体")

        size = r_pr.find(f"{{{W_NS}}}sz")
        size_cs = r_pr.find(f"{{{W_NS}}}szCs")
        self.assertIsNotNone(size)
        self.assertEqual(size.get(f"{{{W_NS}}}val"), "28")
        self.assertIsNotNone(size_cs)
        self.assertEqual(size_cs.get(f"{{{W_NS}}}val"), "28")

    def test_add_bookmark_wraps_target_range(self) -> None:
        from scripts._docx_fields import add_bookmark

        document = docx.Document()
        paragraph = document.add_paragraph("Figure target")

        add_bookmark(paragraph, "fig_0001")

        xml = paragraph._p.xml
        self.assertIn("bookmarkStart", xml)
        self.assertIn("bookmarkEnd", xml)
        self.assertIn('w:name="fig_0001"', xml)
        start_index = xml.index("bookmarkStart")
        text_index = xml.index("Figure target")
        end_index = xml.index("bookmarkEnd")
        self.assertLess(start_index, text_index)
        self.assertLess(text_index, end_index)

    def test_add_complex_field_builds_valid_field_chars(self) -> None:
        from scripts._docx_fields import append_complex_field

        document = docx.Document()
        paragraph = document.add_paragraph()

        append_complex_field(paragraph, ' REF fig_0001 \\h ', display_text="图1")

        xml = paragraph._p.xml
        self.assertIn('fldCharType="begin"', xml)
        self.assertIn('fldCharType="separate"', xml)
        self.assertIn('fldCharType="end"', xml)
        self.assertIn("REF fig_0001", xml)
        self.assertIn(">图1<", xml)

    def test_insert_toc_field_paragraph_creates_standard_toc_code(self) -> None:
        from scripts._docx_fields import insert_toc_field

        document = docx.Document()
        anchor = document.add_paragraph("目录")
        toc_paragraph = insert_paragraph_after(anchor)

        insert_toc_field(toc_paragraph, (1, 3))

        xml = toc_paragraph._p.xml
        self.assertIn("TOC", xml)
        self.assertIn("\\o", xml)
        self.assertNotIn("\\\\o", xml)
        self.assertIn("1-3", xml)

    def test_insert_seq_field_creates_caption_number_field(self) -> None:
        from scripts._docx_fields import append_complex_field

        document = docx.Document()
        paragraph = document.add_paragraph("图")

        append_complex_field(paragraph, " SEQ 图 \\* ARABIC ", display_text="1")

        xml = paragraph._p.xml
        self.assertIn("SEQ 图", xml)
        self.assertIn('fldCharType="begin"', xml)
        self.assertIn('fldCharType="end"', xml)
        self.assertIn("\\* ARABIC", xml)
        self.assertNotIn("\\\\* ARABIC", xml)

    def test_figure_caption_uses_seq_field_and_bookmark(self) -> None:
        project_root = self.create_project()
        image_path = project_root / "docs" / "images" / "arch.png"
        image_path.parent.mkdir(parents=True, exist_ok=True)
        image_path.write_bytes(TEST_PNG_BYTES)
        (project_root / "docs" / "report_body.md").write_text(
            "## Figures\n\n![Architecture](images/arch.png)\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        caption = next(
            paragraph for paragraph in redacted.paragraphs if paragraph.style.name == "图题"
        )

        self.assertIn("SEQ 图", caption._p.xml)
        self.assertIn("bookmarkStart", caption._p.xml)
        self.assertIn("fig_0001", caption._p.xml)
        self.assertIn("\\* ARABIC", caption._p.xml)
        self.assertNotIn("\\\\* ARABIC", caption._p.xml)
        start_index = caption._p.xml.index("bookmarkStart")
        text_index = caption._p.xml.index(">图<")
        end_index = caption._p.xml.index("bookmarkEnd")
        self.assertLess(start_index, text_index)
        self.assertLess(text_index, end_index)

    def test_table_caption_uses_seq_field_and_bookmark(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "## Metrics\n\n| Name | Value |\n| --- | --- |\n| Alpha | 1 |\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        caption = next(
            paragraph for paragraph in redacted.paragraphs if paragraph.style.name == "表题"
        )

        self.assertIn("SEQ 表", caption._p.xml)
        self.assertIn("bookmarkStart", caption._p.xml)
        self.assertIn("tbl_0001", caption._p.xml)
        self.assertIn("\\* ARABIC", caption._p.xml)
        self.assertNotIn("\\\\* ARABIC", caption._p.xml)
        start_index = caption._p.xml.index("bookmarkStart")
        text_index = caption._p.xml.index(">表<")
        end_index = caption._p.xml.index("bookmarkEnd")
        self.assertLess(start_index, text_index)
        self.assertLess(text_index, end_index)

    def test_caption_and_target_use_keep_with_next_not_group_object(self) -> None:
        project_root = self.create_project()
        image_path = project_root / "docs" / "images" / "arch.png"
        image_path.parent.mkdir(parents=True, exist_ok=True)
        image_path.write_bytes(TEST_PNG_BYTES)
        (project_root / "docs" / "report_body.md").write_text(
            "## Figures\n\n![Architecture](images/arch.png)\n\n| Name | Value |\n| --- | --- |\n| Alpha | 1 |\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        figure_paragraph = next(
            paragraph for paragraph in redacted.paragraphs if "<w:drawing" in paragraph._p.xml
        )
        figure_caption = next(
            paragraph for paragraph in redacted.paragraphs if paragraph.style.name == "图题"
        )
        table_caption = next(
            paragraph for paragraph in redacted.paragraphs if paragraph.style.name == "表题"
        )

        self.assertIn("keepNext", figure_paragraph._p.xml)
        self.assertIn("keepLines", figure_paragraph._p.xml)
        self.assertIn("keepNext", figure_caption._p.xml)
        self.assertIn("keepLines", figure_caption._p.xml)
        self.assertIn("keepNext", table_caption._p.xml)
        self.assertNotIn("wpg:wgp", figure_paragraph._p.xml)

    def test_toc_is_inserted_only_when_detected_and_confirmed(self) -> None:
        project_root = self.create_project()
        self.insert_toc_placeholder(project_root)

        self.set_toc_confirmation(
            project_root, enabled=False, needs_confirmation=False
        )
        disabled_result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(disabled_result.returncode, 0, msg=disabled_result.stderr)
        disabled_doc = docx.Document(project_root / "out" / "redacted.docx")
        self.assertFalse(any(" TOC " in paragraph._p.xml for paragraph in disabled_doc.paragraphs))

        self.set_toc_confirmation(
            project_root, enabled=True, needs_confirmation=False
        )
        enabled_result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(enabled_result.returncode, 0, msg=enabled_result.stderr)
        enabled_doc = docx.Document(project_root / "out" / "redacted.docx")
        self.assertTrue(any(" TOC " in paragraph._p.xml for paragraph in enabled_doc.paragraphs))

    def test_build_report_does_not_force_field_update_on_open_by_default(self) -> None:
        project_root = self.create_project()
        image_path = project_root / "docs" / "images" / "arch.png"
        image_path.parent.mkdir(parents=True, exist_ok=True)
        image_path.write_bytes(TEST_PNG_BYTES)
        (project_root / "docs" / "report_body.md").write_text(
            "## Figures\n\n![Architecture](images/arch.png)\n\n[[REF:figure:fig_0001|见下图]]\n",
            encoding="utf-8",
        )
        self.set_cross_reference_confirmation(project_root, True)

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import zipfile

        with zipfile.ZipFile(project_root / "out" / "redacted.docx", "r") as docx_zip:
            settings_xml = docx_zip.read("word/settings.xml").decode("utf-8")
        self.assertNotIn("updateFields", settings_xml)

    def test_inserted_toc_uses_standard_toc_field(self) -> None:
        project_root = self.create_project()
        self.insert_toc_placeholder(project_root)
        self.set_toc_confirmation(project_root, enabled=True, needs_confirmation=False)

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        toc_paragraph = next(
            paragraph for paragraph in redacted.paragraphs if " TOC " in paragraph._p.xml
        )

        self.assertIn('fldCharType="begin"', toc_paragraph._p.xml)
        self.assertIn('fldCharType="separate"', toc_paragraph._p.xml)
        self.assertIn('fldCharType="end"', toc_paragraph._p.xml)
        self.assertIn('TOC \\o "1-3"', toc_paragraph._p.xml)
        self.assertNotIn('TOC \\\\o "1-3"', toc_paragraph._p.xml)

    def test_inserted_toc_uses_separate_title_paragraph_with_title_style(self) -> None:
        project_root = self.create_project()
        self.set_toc_confirmation(project_root, enabled=True, needs_confirmation=False)

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        payload = json.loads(result.stdout)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        toc_title_index = next(
            index
            for index, paragraph in enumerate(redacted.paragraphs)
            if paragraph.text.strip() == "目录"
        )
        toc_field_index = next(
            index
            for index, paragraph in enumerate(redacted.paragraphs)
            if " TOC " in paragraph._p.xml
        )

        self.assertEqual(redacted.paragraphs[toc_title_index].style.name, "题目")
        self.assertLess(toc_title_index, toc_field_index)
        if payload["toc_refresh"]["updated"]:
            self.assertNotEqual(redacted.paragraphs[toc_field_index].text.strip(), "")
            self.assertIn(
                redacted.paragraphs[toc_field_index].style.style_id,
                {"TOC1", "TOC2", "TOC3"},
            )
        else:
            self.assertEqual(redacted.paragraphs[toc_field_index].text.strip(), "")

    @unittest.skipUnless(
        sys.platform == "win32" and WINWORD.exists(),
        "Word automation is only available on Windows with Word installed",
    )
    def test_build_report_populates_toc_entries_when_word_automation_is_available(
        self,
    ) -> None:
        project_root = self.create_project()
        self.set_toc_confirmation(project_root, enabled=True, needs_confirmation=False)

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)

        import zipfile

        with zipfile.ZipFile(project_root / "out" / "redacted.docx", "r") as docx_zip:
            document_xml = docx_zip.read("word/document.xml").decode("utf-8")

        self.assertIn("PAGEREF _Toc", document_xml)
        self.assertIn("_Toc", document_xml)

    def test_build_report_inserts_toc_page_before_body_when_no_placeholder_exists(
        self,
    ) -> None:
        project_root = self.create_project()
        self.set_toc_confirmation(project_root, enabled=True, needs_confirmation=False)

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")

        toc_title_index = next(
            index
            for index, paragraph in enumerate(redacted.paragraphs)
            if paragraph.text.strip() == "目录"
        )
        toc_field_index = next(
            index
            for index, paragraph in enumerate(redacted.paragraphs)
            if " TOC " in paragraph._p.xml
        )
        body_heading_index = next(
            index
            for index, paragraph in enumerate(redacted.paragraphs)
            if paragraph.style.name in {"标题2", "Heading 1"}
        )

        self.assertLess(toc_title_index, toc_field_index)
        self.assertLess(toc_field_index, body_heading_index)
        self.assertIn("w:type=\"page\"", redacted.paragraphs[toc_title_index - 1]._p.xml)
        self.assertTrue(
            any(
                'w:type="page"' in paragraph._p.xml
                for paragraph in redacted.paragraphs[toc_field_index + 1 : body_heading_index]
            )
        )

    def test_build_report_inserts_toc_page_when_enabled_without_detection(
        self,
    ) -> None:
        project_root = self.create_project()
        plan_path = project_root / "config" / "template.plan.json"
        plan = json.loads(plan_path.read_text(encoding="utf-8"))
        toc = plan.setdefault("semantics", {}).setdefault("toc", {})
        toc["detected"] = False
        toc["enabled"] = True
        toc["needs_confirmation"] = False
        toc["kind"] = "none"
        toc["source"] = "manual"
        plan_path.write_text(
            json.dumps(plan, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")

        self.assertTrue(any(" TOC " in paragraph._p.xml for paragraph in redacted.paragraphs))

    def test_build_report_reuses_existing_cover_page_break_for_inserted_toc(
        self,
    ) -> None:
        project_root = self.create_project()
        self.set_toc_confirmation(project_root, enabled=True, needs_confirmation=False)

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")

        body_heading_index = next(
            index
            for index, paragraph in enumerate(redacted.paragraphs)
            if paragraph.style.name in {"标题2", "Heading 1"}
        )
        page_break_count = sum(
            1
            for paragraph in redacted.paragraphs[:body_heading_index]
            if 'w:type="page"' in paragraph._p.xml
        )

        self.assertEqual(page_break_count, 2)

    def test_build_report_collapses_consecutive_cover_page_breaks_around_inserted_toc(
        self,
    ) -> None:
        project_root = self.create_project()
        self.add_extra_cover_page_break(project_root)
        self.set_toc_confirmation(project_root, enabled=True, needs_confirmation=False)

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")

        body_heading_index = next(
            index
            for index, paragraph in enumerate(redacted.paragraphs)
            if paragraph.style.name in {"标题2", "Heading 1"}
        )
        toc_index = next(
            index
            for index, paragraph in enumerate(redacted.paragraphs)
            if " TOC " in paragraph._p.xml
        )
        page_break_indexes = [
            index
            for index, paragraph in enumerate(redacted.paragraphs[:body_heading_index])
            if 'w:type="page"' in paragraph._p.xml
        ]

        self.assertEqual(len(page_break_indexes), 2)
        self.assertGreater(toc_index, page_break_indexes[0])
        self.assertLess(toc_index, page_break_indexes[1])

    def test_default_repo_toc_formatting_matches_policy(self) -> None:
        project_root = self.create_project()
        self.write_report_body(
            project_root,
            "## 一级\n\n### 二级\n\n#### 三级\n\n正文。\n",
        )
        self.set_toc_confirmation(project_root, enabled=True, needs_confirmation=False)

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        payload = json.loads(result.stdout)
        redacted = docx.Document(project_root / "out" / "redacted.docx")

        for style_name in ("目录1", "目录2", "目录3"):
            self.assert_toc_style_matches_policy(redacted.styles[style_name])

        if not payload["toc_refresh"]["updated"]:
            self.skipTest("Word automation unavailable, cannot validate built-in TOC styles")

        toc_entry_styles = {
            paragraph.style.style_id: paragraph.style
            for paragraph in redacted.paragraphs
            if getattr(paragraph.style, "style_id", "") in {"TOC1", "TOC2", "TOC3"}
        }
        self.assertTrue(toc_entry_styles)
        for style in toc_entry_styles.values():
            self.assert_toc_style_matches_policy(style)

    def test_postprocess_cross_reference_replaces_figure_placeholder_with_clickable_label(
        self,
    ) -> None:
        project_root = self.create_project()
        image_path = project_root / "docs" / "images" / "arch.png"
        image_path.parent.mkdir(parents=True, exist_ok=True)
        image_path.write_bytes(TEST_PNG_BYTES)
        (project_root / "docs" / "report_body.md").write_text(
            "## Figures\n\n![Architecture](images/arch.png)\n\n[[REF:figure:fig_0001|见下图]] 展示了系统结构。\n",
            encoding="utf-8",
        )
        self.set_cross_reference_confirmation(project_root, True)

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(
            item for item in redacted.paragraphs if "展示了系统结构" in item.text
        )

        self.assertEqual(paragraph.text.strip(), "见下图1 展示了系统结构。")
        self.assertIn("w:hyperlink", paragraph._p.xml)
        self.assertIn('w:anchor="fig_0001"', paragraph._p.xml)
        self.assertIn(">图1<", paragraph._p.xml)
        self.assertNotIn("[[REF:figure:fig_0001|见下图]]", paragraph._p.xml)

    def test_postprocess_cross_reference_replaces_table_placeholder_with_clickable_label(
        self,
    ) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "## Metrics\n\n| Name | Value |\n| --- | --- |\n| Alpha | 1 |\n\n[[REF:table:tbl_0001|见上表]] 汇总了实验结果。\n",
            encoding="utf-8",
        )
        self.set_cross_reference_confirmation(project_root, True)

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(
            item for item in redacted.paragraphs if "汇总了实验结果" in item.text
        )

        self.assertEqual(paragraph.text.strip(), "见上表1 汇总了实验结果。")
        self.assertIn("w:hyperlink", paragraph._p.xml)
        self.assertIn('w:anchor="tbl_0001"', paragraph._p.xml)
        self.assertIn(">表1<", paragraph._p.xml)
        self.assertNotIn("[[REF:table:tbl_0001|见上表]]", paragraph._p.xml)

    def test_cross_reference_pass_respects_project_confirmation(self) -> None:
        project_root = self.create_project()
        image_path = project_root / "docs" / "images" / "arch.png"
        image_path.parent.mkdir(parents=True, exist_ok=True)
        image_path.write_bytes(TEST_PNG_BYTES)
        (project_root / "docs" / "report_body.md").write_text(
            "## Figures\n\n![Architecture](images/arch.png)\n\n[[REF:figure:fig_0001]]\n",
            encoding="utf-8",
        )
        self.set_cross_reference_confirmation(project_root, False)

        result = self.run_completed(project_root, "build_report.py")
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        redacted = docx.Document(project_root / "out" / "redacted.docx")
        paragraph = next(
            item for item in redacted.paragraphs if "[[REF:figure:fig_0001]]" in item.text
        )

        self.assertEqual(paragraph.text.strip(), "[[REF:figure:fig_0001]]")
        self.assertNotIn("w:hyperlink", paragraph._p.xml)


if __name__ == "__main__":
    unittest.main()
