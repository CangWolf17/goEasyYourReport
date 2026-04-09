from __future__ import annotations

import json
import shutil
import subprocess
import unittest
import uuid
from pathlib import Path
from unittest import mock

import docx
import yaml


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PYTHON = PROJECT_ROOT / ".venv" / "Scripts" / "python.exe"


class ConfirmationPackageTests(unittest.TestCase):
    def test_dependency_files_exist(self) -> None:
        self.assertTrue((PROJECT_ROOT / "pyproject.toml").exists())
        self.assertTrue((PROJECT_ROOT / "requirements.txt").exists())
        self.assertIn("Pygments", (PROJECT_ROOT / "pyproject.toml").read_text())
        self.assertIn("Pygments", (PROJECT_ROOT / "requirements.txt").read_text())

    def test_root_readme_exists_with_repo_specific_content(self) -> None:
        readme_path = PROJECT_ROOT / "README.md"
        self.assertTrue(readme_path.exists())
        readme = readme_path.read_text(encoding="utf-8")
        self.assertIn("uv", readme)
        self.assertIn("build_report.py", readme)
        self.assertIn("preview", readme)

    def test_root_license_is_mit(self) -> None:
        license_path = PROJECT_ROOT / "LICENSE"
        self.assertTrue(license_path.exists())
        license_text = license_path.read_text(encoding="utf-8")
        self.assertIn("Permission is hereby granted, free of charge", license_text)
        self.assertIn("MIT License", license_text)

    def test_repo_docs_describe_docx_integrity_gate_contract(self) -> None:
        skill_text = (PROJECT_ROOT / "SKILL.md").read_text(encoding="utf-8")
        agents_text = (PROJECT_ROOT / "AGENTS.md").read_text(encoding="utf-8")
        readme_text = (PROJECT_ROOT / "README.md").read_text(encoding="utf-8")

        for text in (skill_text, agents_text, readme_text):
            self.assertIn("DOCX integrity gate", text)
            self.assertIn("docx_integrity_error", text)
            self.assertIn("before `verify` or `inject`", text)

    def test_repo_docs_describe_semantic_style_workflow(self) -> None:
        skill_text = (PROJECT_ROOT / "SKILL.md").read_text(encoding="utf-8")
        agents_text = (PROJECT_ROOT / "AGENTS.md").read_text(encoding="utf-8")
        readme_text = (PROJECT_ROOT / "README.md").read_text(encoding="utf-8")

        for text in (skill_text, agents_text, readme_text):
            self.assertIn("semantic template scan", text)
            self.assertIn("style-gap confirmation", text)
            self.assertIn("TOC / reference-block detection in preview", text)
            self.assertIn("semantic style recommendation before build", text)

    def test_repo_docs_describe_toc_and_cross_reference_workflow(self) -> None:
        skill_text = (PROJECT_ROOT / "SKILL.md").read_text(encoding="utf-8")
        agents_text = (PROJECT_ROOT / "AGENTS.md").read_text(encoding="utf-8")
        readme_text = (PROJECT_ROOT / "README.md").read_text(encoding="utf-8")

        for text in (skill_text, agents_text, readme_text):
            self.assertIn("TOC is inserted only when detected and confirmed", text)
            self.assertIn("figure / table cross-references are a post-processing step", text)
            self.assertIn("cross-reference insertion requires user confirmation", text)

    def test_repo_docs_describe_equation_and_bibliography_workflow(self) -> None:
        skill_text = (PROJECT_ROOT / "SKILL.md").read_text(encoding="utf-8")
        agents_text = (PROJECT_ROOT / "AGENTS.md").read_text(encoding="utf-8")
        readme_text = (PROJECT_ROOT / "README.md").read_text(encoding="utf-8")

        for text in (skill_text, agents_text, readme_text):
            self.assertIn("supported equation syntax", text)
            self.assertIn(
                "inline equations render inline, block equations are numbered and cross-referenceable",
                text,
            )
            self.assertIn(
                "bibliography source modes: agent_generate_verified_only, agent_search_and_screen, user_supplied_files",
                text,
            )
            self.assertIn(
                "no reference block in task/template means source-only, not output",
                text,
            )

    def test_init_project_copies_code_theme_sample(self) -> None:
        project_root = self.create_project()

        workflow = json.loads(
            (project_root / "workflow.json").read_text(encoding="utf-8")
        )
        self.assertEqual(workflow["rendering"]["code_blocks"]["theme"], "github-light")
        self.assertEqual(
            workflow["rendering"]["code_blocks"]["theme_override"],
            "./config/code-theme.user.json",
        )
        self.assertTrue(
            (project_root / "config" / "code-theme.user.sample.json").exists()
        )

    def create_project(self) -> Path:
        sandbox_root = PROJECT_ROOT / "temp" / "confirmation-package-tests"
        sandbox_root.mkdir(parents=True, exist_ok=True)
        project_root = sandbox_root / uuid.uuid4().hex
        project_root.mkdir(parents=True, exist_ok=True)
        self.addCleanup(lambda: shutil.rmtree(project_root, ignore_errors=True))
        result = subprocess.run(
            [
                str(PYTHON),
                str(PROJECT_ROOT / "scripts" / "init_project.py"),
                "--project-root",
                str(project_root),
            ],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        return project_root

    def run_json(
        self, project_root: Path, script_name: str, *extra_args: str
    ) -> dict[str, object]:
        result = subprocess.run(
            [
                str(PYTHON),
                str(project_root / "scripts" / script_name),
                "--project-root",
                str(project_root),
                *extra_args,
            ],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, msg=result.stderr)
        return json.loads(result.stdout)

    def run_workflow(
        self, project_root: Path, action: str, *extra_args: str
    ) -> subprocess.CompletedProcess[str]:
        return subprocess.run(
            [
                str(PYTHON),
                str(project_root / "scripts" / "workflow_agent.py"),
                action,
                "--project-root",
                str(project_root),
                *extra_args,
            ],
            capture_output=True,
            text=True,
        )

    def set_ready_to_write(self, project_root: Path, ready: bool = True) -> None:
        task_path = project_root / "report.task.yaml"
        payload = yaml.safe_load(task_path.read_text(encoding="utf-8"))
        payload["task"]["ready_to_write"] = ready
        payload["task"]["needs_user_input"] = not ready
        payload["task"]["stage"] = "ready_to_build" if ready else "collecting_materials"
        payload["runtime"]["next_step"] = "build" if ready else "prepare"
        task_path.write_text(
            yaml.safe_dump(payload, allow_unicode=True, sort_keys=False),
            encoding="utf-8",
        )

    def insert_toc_placeholder(self, project_root: Path) -> None:
        template_path = project_root / "templates" / "template.user.docx"
        template = docx.Document(template_path)
        anchor = next(
            paragraph
            for paragraph in template.paragraphs
            if getattr(paragraph.style, "name", "") == "标题2"
        )
        anchor.insert_paragraph_before("目录")
        template.save(template_path)

        result = subprocess.run(
            [
                str(PYTHON),
                str(project_root / "scripts" / "scan_template.py"),
                "--project-root",
                str(project_root),
            ],
            capture_output=True,
            text=True,
        )
        self.assertEqual(result.returncode, 0, msg=result.stderr)

    def assert_normalized_agent_payload(
        self, payload: dict[str, object], action: str
    ) -> None:
        self.assertEqual(payload["action"], action)
        self.assertIn("status", payload)
        self.assertIn("summary", payload)
        self.assertIn("artifacts", payload)
        self.assertIn("issues", payload)
        self.assertIn("warnings", payload)
        self.assertIn("next_step", payload)
        self.assertIsInstance(payload["issues"], list)
        self.assertIsInstance(payload["warnings"], list)

    def test_scan_template_reports_heading_anchors_and_field_candidates(self) -> None:
        project_root = self.create_project()
        scan = self.run_json(project_root, "scan_template.py")

        anchors = scan["anchors"]
        headings = anchors["headings"]
        field_candidates = anchors["field_candidates"]

        self.assertTrue(any(item["kind"] == "heading" for item in headings))
        self.assertTrue(
            any(str(item["text"]).endswith("：") for item in field_candidates)
        )

    def test_build_preview_writes_summary_and_region_markers(self) -> None:
        project_root = self.create_project()
        preview = self.run_json(project_root, "build_preview.py")

        self.assertIn("summary", preview)
        summary_path = Path(preview["summary"])
        self.assertTrue(summary_path.exists())
        summary = json.loads(summary_path.read_text(encoding="utf-8"))

        self.assertEqual(summary["preview"], "./out/preview.docx")
        self.assertEqual(summary["summary"], "./out/preview.summary.json")
        self.assertTrue(summary["field_binding"]["bindings"])
        self.assertIn("completion_date", summary["field_binding"]["availability"])

        preview_doc = docx.Document(Path(preview["preview"]))
        texts = [
            paragraph.text.strip()
            for paragraph in preview_doc.paragraphs
            if paragraph.text.strip()
        ]

        self.assertTrue(any("Locked Region" in text for text in texts))
        self.assertTrue(any("Fillable Region" in text for text in texts))

    def test_preview_summary_surfaces_report_task_stage(self) -> None:
        project_root = self.create_project()

        result = self.run_workflow(project_root, "prepare")

        self.assertEqual(result.returncode, 0, msg=result.stderr)
        summary = json.loads(
            (project_root / "out" / "preview.summary.json").read_text(
                encoding="utf-8"
            )
        )

        self.assertIn("task_contract", summary)
        self.assertIn("ready_to_write", summary["task_contract"])
        self.assertIn("stage", summary["task_contract"])
        self.assertIn("next_step", summary["task_contract"])

    def test_verify_report_accepts_preview_mode(self) -> None:
        project_root = self.create_project()
        self.run_json(project_root, "build_preview.py")

        result = subprocess.run(
            [
                str(PYTHON),
                str(project_root / "scripts" / "verify_report.py"),
                "--project-root",
                str(project_root),
                "--docx",
                "out/preview.docx",
            ],
            capture_output=True,
            text=True,
        )

        self.assertEqual(result.returncode, 0, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["mode"], "preview")
        self.assertTrue(payload["ok"])

    def test_workflow_agent_prepare_returns_normalized_json(self) -> None:
        project_root = self.create_project()

        result = self.run_workflow(project_root, "prepare")

        self.assertEqual(result.returncode, 0, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assert_normalized_agent_payload(payload, "prepare")
        self.assertEqual(payload["status"], "ok")
        self.assertEqual(payload["issues"], [])
        self.assertTrue(payload["summary"])
        self.assertTrue(payload["artifacts"])
        self.assertTrue(payload["next_step"])

    def test_workflow_agent_build_returns_ok_for_supported_code(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "## Code Example\n\n```python\nprint('ok')\n```",
            encoding="utf-8",
        )
        self.set_ready_to_write(project_root)

        result = self.run_workflow(project_root, "build")

        self.assertEqual(result.returncode, 0, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assert_normalized_agent_payload(payload, "build")
        self.assertEqual(payload["status"], "ok")
        self.assertEqual(payload["issues"], [])
        self.assertEqual(payload["warnings"], [])
        self.assertTrue(payload["summary"])
        self.assertTrue(payload["artifacts"])
        self.assertTrue(payload["next_step"])

    def test_workflow_agent_build_surfaces_unsupported_code_language(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            '## Rust Example\n\n```rust\nfn main() {\n    println!("hi");\n}\n```',
            encoding="utf-8",
        )
        self.set_ready_to_write(project_root)

        result = self.run_workflow(project_root, "build")

        self.assertEqual(result.returncode, 1, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assert_normalized_agent_payload(payload, "build")
        self.assertEqual(payload["status"], "needs_agent_handoff")
        self.assertTrue(payload["issues"])
        self.assertEqual(payload["issues"][0]["kind"], "unsupported_code_language")
        self.assertEqual(payload["issues"][0]["language"], "rust")
        self.assertTrue(payload["artifacts"])

    def test_workflow_agent_build_surfaces_image_insertion_failure(self) -> None:
        project_root = self.create_project()
        (project_root / "docs" / "report_body.md").write_text(
            "## Figures\n\n![Missing](images/missing.png)\n",
            encoding="utf-8",
        )
        self.set_ready_to_write(project_root)

        result = self.run_workflow(project_root, "build")

        self.assertEqual(result.returncode, 1, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assert_normalized_agent_payload(payload, "build")
        self.assertEqual(payload["status"], "needs_agent_handoff")
        self.assertTrue(payload["issues"])
        self.assertEqual(payload["issues"][0]["kind"], "image_insert_failed")
        self.assertEqual(payload["issues"][0]["path"], "images/missing.png")
        self.assertTrue(payload["artifacts"])

    def test_workflow_agent_build_surfaces_docx_integrity_error(self) -> None:
        from scripts import workflow_agent

        project_root = self.create_project()
        build_payload = {
            "redacted": str(project_root / "out" / "redacted.docx"),
            "integrity": {
                "ok": False,
                "errors": [
                    {
                        "kind": "missing_relationship_target",
                        "source": "word/_rels/document.xml.rels",
                        "target": "media/image9.png",
                    }
                ],
            },
            "images": {"inserted": [], "failed": []},
            "code_blocks": {"unsupported": [], "warnings": []},
        }
        script_result = {
            "returncode": 2,
            "stdout": json.dumps(build_payload),
            "stderr": "",
            "json": build_payload,
            "json_error": None,
        }

        with mock.patch(
            "scripts.workflow_agent.run_repo_script",
            return_value=script_result,
        ):
            exit_code, payload = workflow_agent.handle_build(project_root)

        self.assertEqual(exit_code, 2)
        self.assertEqual(payload["status"], "error")
        self.assertEqual(payload["issues"][0]["kind"], "docx_integrity_error")
        self.assertEqual(payload["issues"][0]["details"], build_payload["integrity"]["errors"])

    def test_workflow_agent_preview_surfaces_semantic_confirmation(self) -> None:
        project_root = self.create_project()
        result = self.run_workflow(project_root, "preview")

        self.assertEqual(result.returncode, 1, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assert_normalized_agent_payload(payload, "preview")
        self.assertEqual(payload["status"], "needs_user_confirmation")
        self.assertEqual(payload["next_step"], "review_preview_summary")

    def test_workflow_agent_build_blocks_on_unresolved_toc_confirmation(self) -> None:
        project_root = self.create_project()
        self.insert_toc_placeholder(project_root)
        self.set_ready_to_write(project_root)

        result = self.run_workflow(project_root, "build")

        self.assertEqual(result.returncode, 1, msg=result.stderr)
        payload = json.loads(result.stdout)
        self.assert_normalized_agent_payload(payload, "build")
        self.assertEqual(payload["status"], "needs_user_confirmation")
        self.assertEqual(payload["next_step"], "review_preview_summary")


if __name__ == "__main__":
    unittest.main()
